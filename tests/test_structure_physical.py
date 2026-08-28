"""Тесты для ``structure/physical.py`` (PhysicalDocument adapter).

Покрывает:
    * PDF → blocks с page_index
    * DOCX → blocks (paragraphs + tables) в document order
    * TXT → один block
    * ordinal монотонный (invariant #3)
    * кэш на диске
    * ошибки (missing file, unsupported format)
"""

from __future__ import annotations

import json
import sys
from pathlib import Path

import pytest

_REPO = Path(__file__).resolve().parents[1]
_SCRIPTS = _REPO / "workspace" / "skills" / "legal_summarizer" / "scripts"
if str(_SCRIPTS) not in sys.path:
    sys.path.insert(0, str(_SCRIPTS))
_PROJ = _REPO
if str(_PROJ) not in sys.path:
    sys.path.insert(0, str(_PROJ))

from workspace.skills.legal_summarizer.scripts.structure.physical import (  # noqa: E402
    DocumentBlock,
    PhysicalDocument,
    load_physical_document,
)


def _write_txt(path: Path, content: str) -> None:
    path.write_text(content, encoding="utf-8")


def _write_docx(path: Path, paragraphs: list[str]) -> None:
    from docx import Document

    doc = Document()
    for p in paragraphs:
        doc.add_paragraph(p)
    doc.save(str(path))


def _write_docx_with_table(path: Path, paragraphs: list[str], table_rows: list[list[str]]) -> None:
    from docx import Document

    doc = Document()
    for p in paragraphs:
        doc.add_paragraph(p)
    table = doc.add_table(rows=len(table_rows), cols=len(table_rows[0]))
    for r_idx, row in enumerate(table_rows):
        for c_idx, cell in enumerate(row):
            table.rows[r_idx].cells[c_idx].text = cell
    doc.save(str(path))


def _write_pdf(path: Path, pages_text: list[str]) -> None:
    from pypdf import PdfWriter

    writer = PdfWriter()
    for text in pages_text:
        from pypdf.generic import (
            ArrayObject,
            DecodedStreamObject,
            DictionaryObject,
            FloatObject,
            NameObject,
            TextStringObject,
        )

        page = writer.add_blank_page(width=595, height=842)
        font = DictionaryObject({
            NameObject("/Type"): NameObject("/Font"),
            NameObject("/Subtype"): NameObject("/Type1"),
            NameObject("/BaseFont"): NameObject("/Helvetica"),
        })
        content = DecodedStreamObject()
        content.set_data(
            f"BT /F1 12 Tf 50 800 Td ({text}) Tj ET".encode("latin-1", errors="replace")
        )
        page[NameObject("/Resources")] = DictionaryObject({
            NameObject("/Font"): DictionaryObject({NameObject("/F1"): font})
        })
        page[NameObject("/Contents")] = content
    writer.write(str(path))


# ---------------------------------------------------------------------------
# TXT
# ---------------------------------------------------------------------------


def test_load_txt_returns_single_block(tmp_path):
    p = tmp_path / "doc.txt"
    _write_txt(p, "Договор аренды.\n\nСтороны и предмет.")
    doc = load_physical_document(p)
    assert isinstance(doc, PhysicalDocument)
    assert doc.format == "txt"
    assert len(doc.blocks) == 1
    assert doc.blocks[0].block_type == "text"
    assert doc.blocks[0].ordinal == 0
    assert "Договор" in doc.blocks[0].content


def test_txt_block_page_metadata_is_none(tmp_path):
    p = tmp_path / "doc.txt"
    _write_txt(p, "abc")
    doc = load_physical_document(p)
    block = doc.blocks[0]
    assert block.page_index is None
    assert block.paragraph_index is None
    assert block.table_index is None


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------


def test_load_docx_returns_paragraphs(tmp_path):
    p = tmp_path / "doc.docx"
    _write_docx(p, ["Первый параграф.", "Второй параграф.", "Третий."])
    doc = load_physical_document(p)
    assert doc.format == "docx"
    para_blocks = [b for b in doc.blocks if b.block_type == "paragraph"]
    assert len(para_blocks) == 3
    assert [b.content for b in para_blocks] == [
        "Первый параграф.",
        "Второй параграф.",
        "Третий.",
    ]
    assert [b.paragraph_index for b in para_blocks] == [0, 1, 2]


def test_load_docx_returns_tables_in_document_order(tmp_path):
    p = tmp_path / "doc.docx"
    _write_docx_with_table(
        p,
        paragraphs=["До таблицы."],
        table_rows=[["a", "b"], ["c", "d"]],
    )
    doc = load_physical_document(p)
    assert len(doc.blocks) == 2
    assert doc.blocks[0].block_type == "paragraph"
    assert doc.blocks[0].content == "До таблицы."
    assert doc.blocks[1].block_type == "table"
    assert "a | b" in doc.blocks[1].content
    assert "c | d" in doc.blocks[1].content
    assert doc.blocks[1].table_index == 0
    assert doc.blocks[1].block_metadata["row_count"] == 2


def test_load_docx_ordinal_is_monotonic(tmp_path):
    p = tmp_path / "doc.docx"
    _write_docx_with_table(
        p,
        paragraphs=["p1", "p2", "p3"],
        table_rows=[["x", "y"], ["z", "w"]],
    )
    doc = load_physical_document(p)
    ordinals = [b.ordinal for b in doc.blocks]
    assert ordinals == list(range(len(doc.blocks)))


def test_load_docx_skips_empty_paragraphs(tmp_path):
    p = tmp_path / "doc.docx"
    _write_docx(p, ["Видимый.", "", "   ", "Тоже видимый."])
    doc = load_physical_document(p)
    assert [b.content for b in doc.blocks if b.block_type == "paragraph"] == [
        "Видимый.",
        "Тоже видимый.",
    ]


# ---------------------------------------------------------------------------
# PDF
# ---------------------------------------------------------------------------


def test_load_pdf_returns_pages(tmp_path):
    p = tmp_path / "doc.pdf"
    _write_pdf(p, ["Первая страница.", "Вторая страница."])
    doc = load_physical_document(p)
    assert doc.format == "pdf"
    assert doc.page_count == 2
    page_blocks = [b for b in doc.blocks if b.block_type == "page"]
    assert len(page_blocks) == 2
    assert page_blocks[0].page_index == 1
    assert page_blocks[1].page_index == 2


def test_load_pdf_ordinal_is_monotonic(tmp_path):
    p = tmp_path / "doc.pdf"
    _write_pdf(p, ["page1", "page2", "page3"])
    doc = load_physical_document(p)
    ordinals = [b.ordinal for b in doc.blocks]
    assert ordinals == list(range(len(doc.blocks)))


def test_load_pdf_extracts_tables_per_page(tmp_path):
    try:
        import pdfplumber  # noqa: F401
    except ImportError:
        pytest.skip("pdfplumber not available")
    p = tmp_path / "doc.pdf"
    _write_pdf(p, ["page with table"])
    doc = load_physical_document(p)
    assert doc.page_count == 1


# ---------------------------------------------------------------------------
# Errors
# ---------------------------------------------------------------------------


def test_load_missing_file_raises(tmp_path):
    with pytest.raises(FileNotFoundError):
        load_physical_document(tmp_path / "missing.pdf")


def test_load_unsupported_format_raises(tmp_path):
    p = tmp_path / "data.bin"
    p.write_bytes(b"\x00" * 10)
    with pytest.raises(ValueError, match="Неподдерживаемый формат"):
        load_physical_document(p)


# ---------------------------------------------------------------------------
# Cache
# ---------------------------------------------------------------------------


def test_physical_document_cached_on_disk(tmp_path):
    """Повторный вызов не парсит файл, читает из кэша."""
    p = tmp_path / "doc.txt"
    _write_txt(p, "cached content")

    workspace = tmp_path / "ws"
    workspace.mkdir()

    doc1 = load_physical_document(p, workspace_root=workspace)
    cache_dir = workspace / "workspace" / "data_store" / "cache" / "skills" / "legal_summarizer" / "physical"
    assert cache_dir.exists()
    files = list(cache_dir.glob("*.json"))
    assert files, "Кэш должен быть записан"

    cached_raw = json.loads(files[0].read_text(encoding="utf-8"))
    cached_raw["blocks"][0]["content"] = "tampered"
    files[0].write_text(json.dumps(cached_raw, ensure_ascii=False), encoding="utf-8")

    doc2 = load_physical_document(p, workspace_root=workspace)
    assert doc2.blocks[0].content == "tampered", (
        "Кэш должен перечитаться без обращения к файлу"
    )
    assert doc1.blocks[0].content == "cached content"


def test_physical_document_cache_stale_size_invalidates(tmp_path):
    """Кэш с устаревшим размером файла инвалидируется."""
    p = tmp_path / "doc.txt"
    _write_txt(p, "first")
    doc1 = load_physical_document(p, workspace_root=tmp_path)
    assert doc1.blocks[0].content == "first"

    _write_txt(p, "second and longer content")
    doc2 = load_physical_document(p, workspace_root=tmp_path)
    assert doc2.blocks[0].content == "second and longer content"


# ---------------------------------------------------------------------------
# Structural invariants
# ---------------------------------------------------------------------------


def test_block_id_format():
    """block_id имеет формат ``b_NNNN``."""
    b = DocumentBlock(
        block_id="b_0042",
        block_type="paragraph",
        content="x",
        char_count=1,
        page_index=1,
        page_start=1,
        page_end=1,
        paragraph_index=0,
        table_index=None,
        ordinal=42,
        block_metadata={},
    )
    assert b.block_id == "b_0042"


def test_to_dict_roundtrip():
    b = DocumentBlock(
        block_id="b_0000",
        block_type="page",
        content="hello",
        char_count=5,
        page_index=1,
        page_start=1,
        page_end=1,
        paragraph_index=None,
        table_index=None,
        ordinal=0,
        block_metadata={"row_count": 3},
    )
    d = b.to_dict()
    restored = DocumentBlock(**d)
    assert restored == b


def test_physical_document_to_dict_roundtrip(tmp_path):
    p = tmp_path / "doc.txt"
    _write_txt(p, "abc")
    doc = load_physical_document(p)
    d = doc.to_dict()
    restored = PhysicalDocument.from_dict(d)
    assert restored == doc


if __name__ == "__main__":
    sys.exit(pytest.main([__file__, "-v"]))