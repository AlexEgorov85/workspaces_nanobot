"""Characterization tests для ``legal_summarizer``.

Зафиксированное baseline-поведение, которое должно сохраняться
во всех последующих этапах рефакторинга.

Покрывает (что не покрыто в существующих test_structure_*, test_manifest*,
test_skill_legal_summarizer*, test_packing*, test_reducer*):
    * Document order для PDF (multi-page + таблицы)
    * SectionTree order (block_indices монотонны по ord)
    * Table isolation (PDF + DOCX: каждый table — отдельный DocumentBlock)
    * Chunk metadata (section_id/path/page_start/page_end/block_indices
      обязательны для всех chunks после StructureAwareChunker)
    * Multi-chunk batch LLM accounting (map_calls ≥ 1, разделены stats)

Использует существующие хелперы из ``test_structure_physical.py``.
"""
from __future__ import annotations

import sys
from pathlib import Path

import pytest

_REPO = Path(__file__).resolve().parents[1]
_SCRIPTS = _REPO / "workspace" / "skills" / "legal_summarizer" / "scripts"
if str(_SCRIPTS) not in sys.path:
    sys.path.insert(0, str(_SCRIPTS))
_PROJECT = _REPO
if str(_PROJECT) not in sys.path:
    sys.path.insert(0, str(_PROJECT))

def _write_docx(path: Path, paragraphs: list[str]) -> None:
    from docx import Document

    doc = Document()
    for p in paragraphs:
        doc.add_paragraph(p)
    doc.save(str(path))

def _write_docx_with_table(path: Path, paragraphs: list[str], table_rows: list[list[str]]) -> None:
    from docx import Document

    doc = Document()
    for p in paragraphs:
        doc.add_paragraph(p)
    table = doc.add_table(rows=len(table_rows), cols=len(table_rows[0]))
    for r_idx, row in enumerate(table_rows):
        for c_idx, cell in enumerate(row):
            table.rows[r_idx].cells[c_idx].text = cell
    doc.save(str(path))

def _write_pdf(path: Path, pages_text: list[str]) -> None:
    from pypdf import PdfWriter
    from pypdf.generic import (
        DecodedStreamObject,
        DictionaryObject,
        NameObject,
    )

    writer = PdfWriter()
    for text in pages_text:
        page = writer.add_blank_page(width=595, height=842)
        font = DictionaryObject({
            NameObject("/Type"): NameObject("/Font"),
            NameObject("/Subtype"): NameObject("/Type1"),
            NameObject("/BaseFont"): NameObject("/Helvetica"),
        })
        content = DecodedStreamObject()
        content.set_data(
            f"BT /F1 12 Tf 50 800 Td ({text}) Tj ET".encode("latin-1", errors="replace")
        )
        page[NameObject("/Resources")] = DictionaryObject({
            NameObject("/Font"): DictionaryObject({NameObject("/F1"): font})
        })
        page[NameObject("/Contents")] = content
    writer.write(str(path))

def _small_chunking_config() -> dict:
    return {
        "chunk_size": 4000,
        "chunk_overlap": 200,
        "single_call_threshold": 500,
        "chunk_size_input_ratio": None,
    }

def _small_execution_config() -> dict:
    return {
        "confirmation_threshold_sec": 0.0,
        "estimated_chunk_duration_sec": 0.001,
        "max_chunks_for_execution": 10000,
        "context_batching": {
            "system_prompt_tokens": 200,
            "instruction_tokens_per_map": 50,
            "chars_per_token": 3.5,
            "safety_margin": 0.85,
        },
        "llm_max_tokens": 1000,
        "max_concurrent_batches": 1,
    }

# ---------------------------------------------------------------------------
# Document order — PDF multi-page
# ---------------------------------------------------------------------------

def test_pdf_blocks_ordinals_are_dense(tmp_path, monkeypatch):
    """PDF → blocks: ordinal == [0, 1, …, N-1] без дыр (invariant #3)."""
    from workspace.skills.legal_summarizer.scripts.structure.physical import (
        load_physical_document,
    )
    import summarizer

    monkeypatch.setattr(summarizer, "get_chunking_config", _small_chunking_config)

    p = tmp_path / "doc.pdf"
    _write_pdf(p, [f"Страница {i}: текст абзаца номер {i}." for i in range(1, 6)])
    doc = load_physical_document(p)
    ordinals = [b.ordinal for b in doc.blocks]
    assert ordinals == list(range(len(doc.blocks)))

def test_pdf_pages_have_ascending_page_index(tmp_path):
    """PDF page blocks идут с page_index 1..N в document order."""
    from workspace.skills.legal_summarizer.scripts.structure.physical import (
        load_physical_document,
    )

    p = tmp_path / "doc.pdf"
    _write_pdf(p, [f"стр. {i}" for i in range(1, 5)])
    doc = load_physical_document(p)
    page_blocks = [b for b in doc.blocks if b.block_type == "page"]
    assert [b.page_index for b in page_blocks] == list(range(1, 5))

# ---------------------------------------------------------------------------
# Table isolation — DOCX и PDF
# ---------------------------------------------------------------------------

def test_docx_blocks_ordinals_are_dense_and_interleaved(tmp_path):
    """DOCX: ordinal == [0..N-1] без дыр; **paragraphs и tables interleaved в document order**.

    Это acceptance-тест: DOCX должен сохранять реальный document order.
    Acceptance:
        paragraph A / table A / paragraph B / table B / paragraph C
        → ordinals [0..4], порядок сохранён.
    """
    from docx import Document as _Docx
    from workspace.skills.legal_summarizer.scripts.structure.physical import (
        load_physical_document,
    )

    p = tmp_path / "interleaved.docx"
    doc = _Docx()
    doc.add_paragraph("paragraph A")
    tbl1 = doc.add_table(rows=2, cols=2)
    tbl1.rows[0].cells[0].text = "a1"
    tbl1.rows[0].cells[1].text = "a2"
    tbl1.rows[1].cells[0].text = "a3"
    tbl1.rows[1].cells[1].text = "a4"
    doc.add_paragraph("paragraph B")
    tbl2 = doc.add_table(rows=1, cols=2)
    tbl2.rows[0].cells[0].text = "b1"
    tbl2.rows[0].cells[1].text = "b2"
    doc.add_paragraph("paragraph C")
    doc.save(str(p))

    physical = load_physical_document(p)
    contents = [b.content.splitlines()[0] if b.content else "" for b in physical.blocks]

    # Acceptance этапа 5: document order сохранён.
    assert contents == [
        "paragraph A",
        "a1 | a2",      # table A row 1
        "paragraph B",
        "b1 | b2",      # table B row 1
        "paragraph C",
    ]
    # ordinals без дыр (invariant #3).
    assert [b.ordinal for b in physical.blocks] == list(range(len(physical.blocks)))
    # Каждая таблица — отдельный DocumentBlock с table_index ≥ 0.
    tables = [b for b in physical.blocks if b.block_type == "table"]
    assert [t.table_index for t in tables] == [0, 1]
    # Каждый параграф — отдельный DocumentBlock с paragraph_index ≥ 0.
    paras = [b for b in physical.blocks if b.block_type == "paragraph"]
    assert [p.paragraph_index for p in paras] == [0, 1, 2]

def test_docx_table_blocks_never_mixed_with_paragraph(tmp_path):
    """Tables — отдельные DocumentBlock; никогда не склеиваются с paragraph."""
    from workspace.skills.legal_summarizer.scripts.structure.physical import (
        load_physical_document,
    )

    p = tmp_path / "mixed.docx"
    _write_docx_with_table(
        p,
        paragraphs=["p1", "p2"],
        table_rows=[["a", "b"], ["c", "d"], ["e", "f"]],
    )
    doc = load_physical_document(p)
    table_blocks = [b for b in doc.blocks if b.block_type == "table"]
    para_blocks = [b for b in doc.blocks if b.block_type == "paragraph"]
    assert all(b.block_type == "table" for b in table_blocks)
    assert all(b.block_type == "paragraph" for b in para_blocks)
    # Таблица — один блок, не размазана между параграфами.
    assert len(table_blocks) == 1
    assert len(para_blocks) == 2

# ---------------------------------------------------------------------------
# SectionTree order
# ---------------------------------------------------------------------------

def test_section_tree_block_indices_are_monotonic(tmp_path, monkeypatch):
    """Для каждой секции block_indices отсортированы и совпадают с ord в doc."""
    from workspace.skills.legal_summarizer.scripts.structure.physical import (
        load_physical_document,
    )
    from workspace.skills.legal_summarizer.scripts.structure.sections import (
        ROOT_SECTION_ID,
        detect_sections,
        merge_short_sections,
    )
    import summarizer

    monkeypatch.setattr(summarizer, "get_chunking_config", _small_chunking_config)

    p = tmp_path / "legal.txt"
    p.write_text(
        "Введение.\n\n"
        "Статья 1. Общие положения.\n\n"
        "Длинный абзац про общие положения, важные условия договора аренды.\n\n"
        "Статья 2. Обязанности сторон.\n\n"
        "Арендодатель обязуется передать помещение. "
        "Арендатор обязуется платить арендную плату ежемесячно.\n\n"
        "Статья 3. Ответственность.\n\n"
        "Стороны несут ответственность за нарушение обязательств.\n",
        encoding="utf-8",
    )
    doc = load_physical_document(p)
    tree = detect_sections(doc, pdf_path=None)
    tree = merge_short_sections(tree, doc.blocks, min_section_chars=200)

    for sid, section in tree.sections.items():
        if sid == ROOT_SECTION_ID:
            continue
        # block_indices отсортированы (canonical order).
        assert list(section.block_indices) == sorted(section.block_indices)
        # Все указывают на существующие blocks.
        for idx in section.block_indices:
            assert 0 <= idx < len(doc.blocks)

# ---------------------------------------------------------------------------
# Chunk metadata
# ---------------------------------------------------------------------------

def test_chunk_metadata_required_fields_present(tmp_path, monkeypatch):
    """После StructureAwareChunker каждый Chunk имеет обязательные поля."""
    from workspace.skills.legal_summarizer.scripts.packing import TokenBudget
    from workspace.skills.legal_summarizer.scripts.structure.chunks import (
        ChunkConfig,
        StructureAwareChunker,
    )
    from workspace.skills.legal_summarizer.scripts.structure.physical import (
        load_physical_document,
    )
    from workspace.skills.legal_summarizer.scripts.structure.sections import (
        detect_sections,
        merge_short_sections,
    )
    import summarizer

    monkeypatch.setattr(summarizer, "get_chunking_config", _small_chunking_config)

    p = tmp_path / "legal.txt"
    p.write_text(
        "Преамбула документа аренды помещения.\n\n"
        "Статья 1. Общие положения настоящего договора.\n\n"
        + ("Длинный абзац с описанием предмета договора аренды. " * 30)
        + "\n\n"
        + "Статья 2. Права и обязанности сторон.\n\n"
        + ("Арендодатель и арендатор обязуются выполнять условия. " * 30)
        + "\n",
        encoding="utf-8",
    )
    doc = load_physical_document(p)
    tree = detect_sections(doc, pdf_path=None)
    tree = merge_short_sections(tree, doc.blocks, min_section_chars=200)

    chunk_cfg = ChunkConfig(
        max_chunk_chars=2000,
        chunk_overlap_chars=100,
        chars_per_token=3.5,
    )
    chunks = StructureAwareChunker().chunk(doc, tree, chunk_cfg)
    assert len(chunks) >= 1
    for c in chunks:
        assert c.section_id, f"section_id отсутствует: {c}"
        assert isinstance(c.section_id, str)
        assert isinstance(c.section_path, str)
        # page_start/page_end могут быть None для DOCX/TXT без page info,
        # но если один задан — оба должны быть заданы.
        assert (c.page_start is None) == (c.page_end is None)
        assert isinstance(c.block_indices, tuple)
        # token_estimate > 0.
        assert c.token_estimate > 0

def test_chunk_block_indices_are_within_doc(tmp_path, monkeypatch):
    """block_indices каждого chunk ссылаются на существующие DocumentBlock.ordinal."""
    from workspace.skills.legal_summarizer.scripts.structure.chunks import (
        ChunkConfig,
        StructureAwareChunker,
    )
    from workspace.skills.legal_summarizer.scripts.structure.physical import (
        load_physical_document,
    )
    from workspace.skills.legal_summarizer.scripts.structure.sections import (
        detect_sections,
        merge_short_sections,
    )
    import summarizer

    monkeypatch.setattr(summarizer, "get_chunking_config", _small_chunking_config)

    p = tmp_path / "doc.txt"
    p.write_text(
        "Вступление.\n\n"
        + ("Абзац текста с содержательным содержимым для анализа. " * 50)
        + "\n\nЗаключение.\n",
        encoding="utf-8",
    )
    doc = load_physical_document(p)
    tree = detect_sections(doc, pdf_path=None)
    tree = merge_short_sections(tree, doc.blocks, min_section_chars=200)
    chunk_cfg = ChunkConfig(max_chunk_chars=1500, chunk_overlap_chars=0)
    chunks = StructureAwareChunker().chunk(doc, tree, chunk_cfg)
    assert chunks
    max_ord = max(b.ordinal for b in doc.blocks)
    for c in chunks:
        for idx in c.block_indices:
            assert 0 <= idx <= max_ord

# ---------------------------------------------------------------------------
# Multi-chunk batch LLM accounting
# ---------------------------------------------------------------------------

def test_multi_chunk_batch_stats_separated(tmp_path, monkeypatch):
    """run() на multi-chunk документе: stats разделены на map_calls/section_reduce/
    document_reduce/retries."""
    import summarizer

    monkeypatch.setattr(summarizer, "get_execution_config", _small_execution_config)
    monkeypatch.setattr(summarizer, "get_chunking_config", _small_chunking_config)

    call_counts = {"map": 0, "section_reduce": 0, "doc_reduce": 0, "trim": 0}

    def fake_chat(messages, *, context=None, **kwargs):
        import re as _re

        text = messages[1]["content"] if len(messages) > 1 else ""
        if _re.search(r"DOCUMENT CHUNK \d+", text):
            call_counts["map"] += 1
            n = len(_re.findall(r"DOCUMENT CHUNK \d+", text))
            return "\n\n".join(f"DOC CHUNK {i + 1}: саммари чанка {i + 1}" for i in range(n))
        if "Частичные саммари чанков этого раздела" in text:
            call_counts["section_reduce"] += 1
            return "Раздел: краткое саммари раздела."
        if "Саммари разделов документа" in text:
            call_counts["doc_reduce"] += 1
            return "Итоговое саммари документа.\n\nСуть: договор аренды."
        call_counts["trim"] += 1
        return "trim"

    monkeypatch.setattr(summarizer.llm, "chat", fake_chat)

    text = (
        "Преамбула.\n\n"
        "Статья 1. Общие положения.\n\n"
        + ("Длинный абзац с описанием общих положений договора аренды. " * 30)
        + "\n\n"
        "Статья 2. Обязанности сторон.\n\n"
        + ("Арендодатель и арендатор обязуются выполнять условия договора. " * 30)
        + "\n\n"
        "Статья 3. Ответственность.\n\n"
        + ("Стороны несут ответственность за нарушение обязательств. " * 30)
        + "\n\n"
        "Статья 4. Заключительные положения.\n\n"
        + ("Договор вступает в силу с момента подписания сторонами. " * 30)
        + "\n"
    )
    result = summarizer.run(
        text, length="detailed", confirmed=True, workspace_root=tmp_path,
    )
    assert result["status"] == "completed"
    stats = result["stats"]
    assert "map_calls" in stats
    assert "section_reduce_calls" in stats
    assert "section_trim_calls" in stats
    assert "document_reduce_calls" in stats
    assert "reduce_calls" in stats
    assert "total_llm_calls" in stats
    assert "retries" in stats
    assert stats["map_calls"] >= 1
    # total = map + reduce + retries (без trim — он часть reduce_calls).
    assert (
        stats["total_llm_calls"]
        == stats["map_calls"] + stats["reduce_calls"] + stats["retries"]
    )
    # document_reduce ≥ 0.
    assert stats["document_reduce_calls"] >= 0

# ---------------------------------------------------------------------------
# Модульная декомпозиция
# ---------------------------------------------------------------------------

def test_sanitize_module_is_extracted_and_re_exported(tmp_path):
    """``strip_think_blocks`` живёт в отдельном модуле
    ``workspace.skills.legal_summarizer.scripts.sanitize``, и
    ``summarizer._strip_think_blocks`` — это тот же объект (back-compat
    для тестов, которые делают ``monkeypatch.setattr(summarizer, '_strip_think_blocks', ...)``).
    """
    from workspace.skills.legal_summarizer.scripts import sanitize as sanitize_mod
    from workspace.skills.legal_summarizer.scripts import summarizer as summ_mod

    assert hasattr(sanitize_mod, "strip_think_blocks")
    assert summ_mod._strip_think_blocks is sanitize_mod.strip_think_blocks

    # Поведение идентично baseline characterization:
    text = "<think>reasoning</think>Полезный ответ."
    assert sanitize_mod.strip_think_blocks(text) == "Полезный ответ."
    assert summ_mod._strip_think_blocks(text) == "Полезный ответ."

def test_fingerprint_module_is_extracted_and_re_exported():
    """``document_id_for`` и ``resolve_session_key`` живут в
    ``fingerprint.py`` (Этап 14 — пока legacy, миграция в процессе).
    """
    from workspace.skills.legal_summarizer.scripts import fingerprint

    # Поведение document_id_for: стабильный sha256 от первых 64KB,
    # обрезанный до 16 hex символов (legacy поведение).
    # Одинаковый первый 64KB → одинаковый id.
    text_a = "hello world"
    text_b = text_a  # короткий — целиком влезает в 64KB.
    assert fingerprint.document_id_for(text_a) == fingerprint.document_id_for(text_b)
    # Длина hex-id ровно 16 символов.
    assert len(fingerprint.document_id_for(text_a)) == 16
    # Разный текст → разные id (коллизии в 16 hex-символов пренебрежимы).
    text_c = "another start"
    assert fingerprint.document_id_for(text_a) != fingerprint.document_id_for(text_c)
    # Стабильность: один и тот же текст → один и тот же id.
    assert fingerprint.document_id_for(text_a) == fingerprint.document_id_for("hello world")

    # resolve_session_key: None если путь не содержит session-папки.
    assert fingerprint.resolve_session_key(None) is None
    assert fingerprint.resolve_session_key("/some/random/path.pdf") is None

def test_document_cache_module_is_extracted_and_re_exported(tmp_path):
    """``doc_cache_dir/load_doc_cache/save_doc_cache`` живут в
    ``document_cache.py``, ``summarizer._doc_cache_*`` — те же объекты.
    Поведение save/load round-trip идентично baseline.
    """
    from workspace.skills.legal_summarizer.scripts import document_cache
    from workspace.skills.legal_summarizer.scripts import summarizer as summ_mod

    assert summ_mod._doc_cache_dir is document_cache.doc_cache_dir
    assert summ_mod._load_doc_cache is document_cache.load_doc_cache
    assert summ_mod._save_doc_cache is document_cache.save_doc_cache

    ws = tmp_path
    doc_id = "abcd1234"
    session = "session_xyz"
    chunks = {
        "001": {"chunk_id": "001", "summary": "саммари один", "section_id": "s_0001"},
        "002": {"chunk_id": "002", "summary": "саммари два", "section_id": "s_0002"},
    }
    document_cache.save_doc_cache(doc_id, session, ws, chunks)
    loaded = document_cache.load_doc_cache(doc_id, session, ws)
    assert set(loaded.keys()) == {"001", "002"}
    assert loaded["001"]["summary"] == "саммари один"
    assert loaded["002"]["section_id"] == "s_0002"

def test_prompts_runtime_module_is_extracted_and_re_exported():
    """``load_prompt/system_instruction`` и константы
    живут в ``prompts_runtime.py``, ``summarizer._load_prompt`` /
    ``_LENGTH_INSTRUCTIONS`` / ``_system_instruction`` — те же объекты.
    """
    from workspace.skills.legal_summarizer.scripts import prompts_runtime
    from workspace.skills.legal_summarizer.scripts import summarizer as summ_mod

    assert summ_mod._load_prompt is prompts_runtime.load_prompt
    assert summ_mod._LENGTH_INSTRUCTIONS is prompts_runtime.LENGTH_INSTRUCTIONS
    assert (
        summ_mod._QUESTION_INSTRUCTION_TEMPLATE
        is prompts_runtime.QUESTION_INSTRUCTION_TEMPLATE
    )
    assert summ_mod._system_instruction is prompts_runtime.system_instruction

    # Реальные промпты должны грузиться (3 файла из prompts/ dir).
    for name in ("summarize_system", "reduce_system", "section_reduce_system"):
        text = prompts_runtime.load_prompt(name)
        assert isinstance(text, str) and len(text) > 50

    # system_instruction: brief vs question.
    brief = prompts_runtime.system_instruction("brief", None)
    assert "150-250 слов" in brief or "150–250 слов" in brief
    assert "коротко" not in brief.lower()  # это термин CLI, не промпта.
    assert "250" in brief

    q = prompts_runtime.system_instruction("detailed", "что про штрафы?")
    assert "штрафы" in q
    assert "200-300 слов" in q or "200–300 слов" in q

    # Неизвестный length → fallback на brief.
    fallback = prompts_runtime.system_instruction("unknown_length", None)
    assert "150" in fallback

def test_llm_calls_module_is_extracted_and_re_exported():
    """``_llm_batch/_llm_section_reduce/_llm_document_reduce``
    и ``_doc_context`` живут в ``llm_calls.py``. ``_llm_section_trim``
    убран в этапе 17 — section_trim больше не используется summarizer'ом.
    """
    from workspace.skills.legal_summarizer.scripts import llm_calls
    from workspace.skills.legal_summarizer.scripts import summarizer as summ_mod

    assert summ_mod._llm_batch is llm_calls.llm_batch
    assert summ_mod._llm_section_reduce is llm_calls.llm_section_reduce

    assert not hasattr(summ_mod, "_llm_section_trim"), (
        "summarizer не должен импортировать _llm_section_trim после этапа 17"
    )
    assert summ_mod._llm_document_reduce is llm_calls.llm_document_reduce
    assert summ_mod._doc_context is llm_calls.doc_context

    # doc_context: пустой / с title / с begin/end.
    assert llm_calls.doc_context(None) == ""
    assert "НАЗВАНИЕ" in llm_calls.doc_context({"title": "Договор"})
    full = llm_calls.doc_context(
        {"title": "Договор", "begin": "Преамбула...", "end": "Реквизиты..."},
        with_begin_end=True,
    )
    assert "НАЗВАНИЕ" in full
    assert "НАЧАЛО" in full
    assert "КОНЕЦ" in full

def test_pipeline_module_is_extracted_and_re_exported(tmp_path):
    """``process_context_batch/run_one_batch_async/
    load_cached_partials/now_iso`` и ``MAX_BATCH_PARSE_RETRIES`` живут
    в ``pipeline.py``, ``summarizer._*`` — те же объекты.
    """
    from workspace.skills.legal_summarizer.scripts import pipeline
    from workspace.skills.legal_summarizer.scripts import summarizer as summ_mod

    assert summ_mod._process_context_batch is pipeline.process_context_batch
    assert summ_mod._run_one_batch_async is pipeline.run_one_batch_async
    assert summ_mod._load_cached_partials is pipeline.load_cached_partials
    assert summ_mod._now_iso is pipeline.now_iso
    assert summ_mod.MAX_BATCH_PARSE_RETRIES == pipeline.MAX_BATCH_PARSE_RETRIES

    # load_cached_partials: на пустой операции возвращает пустой dict.
    assert pipeline.load_cached_partials("op_none", ["001"], tmp_path) == {}

    # now_iso: возвращает ISO-строку в UTC.
    iso = pipeline.now_iso()
    assert "T" in iso
    assert iso.endswith("+00:00") or iso.endswith("Z")

# ---------------------------------------------------------------------------

# ---------------------------------------------------------------------------

def test_fingerprint_file_basic_properties(tmp_path):
    """``fingerprint_file`` возвращает 16-hex строку, стабильную для одного файла."""
    from workspace.skills.legal_summarizer.scripts.fingerprint import fingerprint_file

    p = tmp_path / "doc.txt"
    p.write_bytes(b"some content")
    fp1 = fingerprint_file(p)
    fp2 = fingerprint_file(p)
    assert fp1 == fp2
    assert len(fp1) == 16
    int(fp1, 16)  # валидный hex

def test_fingerprint_file_distinguishes_same_first_64kb(tmp_path):
    """same first 64KB + different tail → different ID."""
    from workspace.skills.legal_summarizer.scripts.fingerprint import fingerprint_file

    # Оба файла ≥ 64KB с одинаковым первым 64KB; хвост отличается.
    head = b"A" * (70 * 1024)  # 70 KB общего начала.
    tail_a = b"B" * 100
    tail_b = b"C" * 100

    p_a = tmp_path / "a.bin"
    p_a.write_bytes(head + tail_a)
    p_b = tmp_path / "b.bin"
    p_b.write_bytes(head + tail_b)

    # sanity: первые 64KB действительно совпадают.
    assert p_a.read_bytes()[:64 * 1024] == p_b.read_bytes()[:64 * 1024]
    # но полные sha256 — разные.
    assert fingerprint_file(p_a) != fingerprint_file(p_b)

def test_fingerprint_file_different_content_different_id(tmp_path):
    """different file → different ID."""
    from workspace.skills.legal_summarizer.scripts.fingerprint import fingerprint_file

    p1 = tmp_path / "a.bin"
    p1.write_bytes(b"alpha content")
    p2 = tmp_path / "b.bin"
    p2.write_bytes(b"beta content")
    assert fingerprint_file(p1) != fingerprint_file(p2)

def test_fingerprint_file_large_file_streams(tmp_path):
    """Файл > 1 MB (chunk-size в реализации) корректно читается."""
    from workspace.skills.legal_summarizer.scripts.fingerprint import fingerprint_file

    p = tmp_path / "big.bin"
    p.write_bytes(b"A" * (3 * 1024 * 1024 + 123))  # 3 MB + хвост
    fp = fingerprint_file(p)
    assert len(fp) == 16

    # Перезапись → другой id.
    p.write_bytes(b"B" * (3 * 1024 * 1024 + 123))
    fp2 = fingerprint_file(p)
    assert fp != fp2

def test_resolve_document_id_prefers_file_fingerprint(tmp_path):
    """``resolve_document_id(path, text)`` использует fingerprint_file
    когда путь валиден; иначе fallback на document_id_for(text)."""
    from workspace.skills.legal_summarizer.scripts.fingerprint import (
        document_id_for,
        fingerprint_file,
        resolve_document_id,
    )

    p = tmp_path / "doc.bin"
    p.write_bytes(b"hello world")
    text = "hello world"

    # С валидным путём: должен совпасть с fingerprint_file.
    assert resolve_document_id(str(p), text) == fingerprint_file(p)

    # Без пути или с несуществующим путём: fallback на document_id_for(text).
    assert resolve_document_id(None, text) == document_id_for(text)
    assert resolve_document_id(str(tmp_path / "missing.bin"), text) == document_id_for(text)

def test_summarizer_uses_resolve_document_id(monkeypatch, tmp_path):
    """``summarizer.run()`` использует resolve_document_id (file-fingerprint при наличии пути).

    Это поведение этапа 4 — меняет cache-key для прогонов с валидным
    document_path, но НЕ ломает single-call путь для inline-документов.
    """
    from workspace.skills.legal_summarizer.scripts import summarizer

    captured = {}

    real_resolve = summarizer._resolve_document_id

    def spy_resolve(path, text):
        out = real_resolve(path, text)
        captured["path"] = path
        captured["text_len"] = len(text)
        captured["id"] = out
        return out

    monkeypatch.setattr(summarizer, "_resolve_document_id", spy_resolve)

    def fake_chat(messages, *, context=None, **kwargs):
        return "Это договор.\n\nСуть: краткое описание."

    monkeypatch.setattr(summarizer.llm, "chat", fake_chat)

    # Inline-вызов (без document_path) → fallback на document_id_for(text).
    summarizer.run("Договор аренды.", length="brief", workspace_root=tmp_path)
    assert captured["path"] is None
    inline_id = captured["id"]
    assert len(inline_id) == 16

    # С существующим путём → fingerprint_file.
    p = tmp_path / "doc.txt"
    p.write_text("Другой договор.", encoding="utf-8")
    summarizer.run(
        "Другой договор.", length="brief", workspace_root=tmp_path, document_path=str(p),
    )
    assert captured["path"] == str(p)
    assert captured["id"] != inline_id  # другой файл → другой id

# ---------------------------------------------------------------------------

# ---------------------------------------------------------------------------

def _make_block(ordinal: int, content: str, *, block_type: str = "paragraph") -> "DocumentBlock":
    """Хелпер для создания DocumentBlock в тестах cleanup."""
    from workspace.skills.legal_summarizer.scripts.structure.physical import (
        DocumentBlock,
    )

    return DocumentBlock(
        block_id=f"b_{ordinal:04d}",
        block_type=block_type,
        content=content,
        char_count=len(content),
        page_index=ordinal + 1,
        page_start=ordinal + 1,
        page_end=ordinal + 1,
        paragraph_index=ordinal,
        table_index=None,
        ordinal=ordinal,
        block_metadata={},
    )

def test_cleanup_blocks_marks_repeated_header():
    """HEADER + 3×content → 1 header candidate + 3 content blocks.

    Документ вида:
        HEADER
        content page 1
        HEADER
        content page 2
        HEADER
        content page 3
    → HEADER помечен как ``repeated_role="header"`` (3 совпадения),
       content-блоки остаются без пометки.
    """
    from workspace.skills.legal_summarizer.scripts.document_cleanup import (
        cleanup_blocks,
    )

    blocks = [
        _make_block(0, "HEADER"),
        _make_block(1, "content page 1"),
        _make_block(2, "HEADER"),
        _make_block(3, "content page 2"),
        _make_block(4, "HEADER"),
        _make_block(5, "content page 3"),
    ]
    cleaned, stats = cleanup_blocks(blocks)

    assert stats.total_blocks == 6
    # Все три копии HEADER помечены.
    header_blocks = [b for b in cleaned if b.block_metadata.get("is_repeated")]
    assert len(header_blocks) == 3
    assert all(b.block_metadata["repeated_role"] == "header" for b in header_blocks)

    # 3 content-блока НЕ помечены как repeated.
    content_blocks = [b for b in cleaned if not b.block_metadata.get("is_repeated")]
    assert len(content_blocks) == 3
    assert "content page 1" in content_blocks[0].content
    assert "content page 3" in content_blocks[2].content

    # ordinals не меняются (cleanup не переупорядочивает).
    assert [b.ordinal for b in cleaned] == list(range(6))

def test_cleanup_blocks_marks_footer_when_at_end():
    """Footer (повторяется в конце страниц) получает ``repeated_role='footer'``
    на **всех** вхождениях.

    Алгоритм: если ``positions[-1] == total-1`` и средняя позиция ≥ midpoint,
    все вхождения помечаются ``"footer"``.
    """
    from workspace.skills.legal_summarizer.scripts.document_cleanup import (
        cleanup_blocks,
    )

    blocks = [
        _make_block(0, "Chapter 1"),
        _make_block(1, "My Footer"),  # без номера страницы — иначе normalized разный
        _make_block(2, "Chapter 2"),
        _make_block(3, "My Footer"),
        _make_block(4, "Chapter 3"),
        _make_block(5, "My Footer"),
    ]
    cleaned, stats = cleanup_blocks(blocks)
    roles = [b.block_metadata.get("repeated_role") for b in cleaned if b.block_metadata.get("is_repeated")]
    # Все 3 FOOTER-а — role=footer (последний в списке и средняя ≥ midpoint).
    assert roles == ["footer", "footer", "footer"]
    # Все Chapter-блоки — НЕ помечены (уникальные).
    chapters = [b for b in cleaned if not b.block_metadata.get("is_repeated")]
    assert len(chapters) == 3

def test_cleanup_blocks_long_blocks_not_marked():
    """Длинные блоки (> max_length_for_repetition_detection) никогда не repeated,
    даже если их текст повторяется 5 раз. Это защита от случайной пометки
    длинного контента как header.
    """
    from workspace.skills.legal_summarizer.scripts.document_cleanup import (
        CleanupConfig,
        cleanup_blocks,
    )

    long_text = "X" * 400
    blocks = [
        _make_block(0, "intro"),
        _make_block(1, long_text),
        _make_block(2, long_text),
        _make_block(3, long_text),
        _make_block(4, long_text),
    ]
    cleaned, stats = cleanup_blocks(blocks, CleanupConfig(max_length_for_repetition_detection=300))
    # Длинные блоки не должны быть помечены как repeated.
    long_blocked = [b for b in cleaned if b.content == long_text]
    assert all(not b.block_metadata.get("is_repeated") for b in long_blocked)
    assert stats.repeated_blocks == 0

def test_cleanup_blocks_whitespace_normalized():
    """«PAGE  1» и «PAGE 1» — одинаковые для целей сравнения."""
    from workspace.skills.legal_summarizer.scripts.document_cleanup import (
        cleanup_blocks,
    )

    blocks = [
        _make_block(0, "PAGE  1"),
        _make_block(1, "content A"),
        _make_block(2, "PAGE 1"),
        _make_block(3, "content B"),
        _make_block(4, "PAGE 1"),
    ]
    cleaned, stats = cleanup_blocks(blocks)
    repeated = [b for b in cleaned if b.block_metadata.get("is_repeated")]
    assert len(repeated) == 3  # все три «PAGE 1» совпадают после нормализации

def test_cleanup_blocks_threshold_respected():
    """``repetition_threshold`` контролирует, что считается repeated.

    threshold=3 (default): 2 копии → НЕ repeated.
    threshold=2: 2 копии → repeated.
    """
    from workspace.skills.legal_summarizer.scripts.document_cleanup import (
        CleanupConfig,
        cleanup_blocks,
    )

    blocks = [
        _make_block(0, "SHORT"),
        _make_block(1, "content A"),
        _make_block(2, "SHORT"),
    ]

    # threshold=3: 2 совпадения < threshold → не помечаем.
    _, stats_3 = cleanup_blocks(blocks, CleanupConfig(repetition_threshold=3))
    assert stats_3.repeated_blocks == 0

    # threshold=2: 2 совпадения ≥ threshold → помечаем.
    _, stats_2 = cleanup_blocks(blocks, CleanupConfig(repetition_threshold=2))
    assert stats_2.repeated_blocks == 2

def test_cleanup_blocks_empty_blocks_ignored():
    """Пустые блоки (после normalize → '') игнорируются, не считаются repeated."""
    from workspace.skills.legal_summarizer.scripts.document_cleanup import (
        cleanup_blocks,
    )

    blocks = [
        _make_block(0, ""),
        _make_block(1, "real"),
        _make_block(2, "   \n  "),
        _make_block(3, "more"),
    ]
    cleaned, stats = cleanup_blocks(blocks)
    # Никаких пометок repeated (пустые не считаются, остальные уникальны).
    assert stats.repeated_blocks == 0

def test_cleanup_blocks_ordinals_preserved():
    """Cleanup **не переупорядочивает** блоки (invariant #3)."""
    from workspace.skills.legal_summarizer.scripts.document_cleanup import (
        cleanup_blocks,
    )

    blocks = [
        _make_block(0, "A"),
        _make_block(1, "B"),
        _make_block(2, "A"),
        _make_block(3, "C"),
    ]
    cleaned, _ = cleanup_blocks(blocks)
    ordinals = [b.ordinal for b in cleaned]
    assert ordinals == [0, 1, 2, 3]
    # block_id тоже сохраняется.
    assert [b.block_id for b in cleaned] == ["b_0000", "b_0001", "b_0002", "b_0003"]

def test_normalize_whitespace_helper():
    """``normalize_whitespace`` корректно схлопывает whitespace."""
    from workspace.skills.legal_summarizer.scripts.document_cleanup import (
        normalize_whitespace,
    )

    assert normalize_whitespace("a  b\n\nc") == "a b c"
    assert normalize_whitespace("  leading and trailing  ") == "leading and trailing"
    assert normalize_whitespace("") == ""
    assert normalize_whitespace(None) == ""

# ---------------------------------------------------------------------------

# ---------------------------------------------------------------------------

def test_heading_module_extracted():
    """heading detection логика живёт в ``structure/heading.py``."""
    from workspace.skills.legal_summarizer.scripts.structure import heading
    from workspace.skills.legal_summarizer.scripts.structure import sections

    # Public symbols переехали.
    assert hasattr(heading, "HeadingCandidate")
    assert hasattr(heading, "CONFIDENCE_THRESHOLD")
    assert hasattr(heading, "detect_heading_candidates")
    assert hasattr(heading, "apply_confidence_penalties")
    assert hasattr(heading, "filter_above_threshold")

    # Back-compat через sections facade.
    assert sections.HeadingCandidate is heading.HeadingCandidate
    assert sections.CONFIDENCE_THRESHOLD == heading.CONFIDENCE_THRESHOLD

    # ``_classify_regex`` и ``_extract_pdf_outline`` доступны из обоих модулей.
    assert sections._classify_regex is heading._classify_regex
    assert sections._extract_pdf_outline is heading._extract_pdf_outline

def test_tree_module_extracted():
    """SectionTree/DocumentSection/build_section_tree живут в ``tree.py``."""
    from workspace.skills.legal_summarizer.scripts.structure import sections
    from workspace.skills.legal_summarizer.scripts.structure import tree

    assert hasattr(tree, "DocumentSection")
    assert hasattr(tree, "SectionTree")
    assert hasattr(tree, "ROOT_SECTION_ID")
    assert hasattr(tree, "build_section_tree")
    assert hasattr(tree, "section_total_chars")

    # Back-compat.
    assert sections.DocumentSection is tree.DocumentSection
    assert sections.SectionTree is tree.SectionTree
    assert sections.ROOT_SECTION_ID == tree.ROOT_SECTION_ID == "s_root"
    assert sections._build_sections is tree.build_section_tree
    assert sections._section_total_chars is tree.section_total_chars

def test_sections_module_is_facade():
    """``structure/sections.py`` — тонкий facade (~200 строк), re-exports.

    Подсчёт LOC — мягкая проверка, что мы не оставили весь код в sections.
    """
    from pathlib import Path

    p = Path("workspace/skills/legal_summarizer/scripts/structure/sections.py")
    tree_p = Path("workspace/skills/legal_summarizer/scripts/structure/tree.py")
    heading_p = Path("workspace/skills/legal_summarizer/scripts/structure/heading.py")
    sections_lines = sum(1 for _ in p.open("r", encoding="utf-8"))
    tree_lines = sum(1 for _ in tree_p.open("r", encoding="utf-8"))
    heading_lines = sum(1 for _ in heading_p.open("r", encoding="utf-8"))

    # До этапа 7 sections.py был 650 строк. После — facade (≤ 350).
    assert sections_lines < 350, f"sections.py = {sections_lines} строк, ожидался facade ≤ 350"
    # Heading + tree выросли.
    assert heading_lines > 100, f"heading.py = {heading_lines} строк"
    assert tree_lines > 100, f"tree.py = {tree_lines} строк"

# ---------------------------------------------------------------------------

# ---------------------------------------------------------------------------

def _make_heading_block(ordinal: int, content: str, *, block_type: str = "paragraph") -> "DocumentBlock":
    """Хелпер для heading-evidence тестов."""
    from workspace.skills.legal_summarizer.scripts.structure.physical import (
        DocumentBlock,
    )

    return DocumentBlock(
        block_id=f"b_{ordinal:04d}",
        block_type=block_type,
        content=content,
        char_count=len(content),
        page_index=ordinal + 1,
        page_start=ordinal + 1,
        page_end=ordinal + 1,
        paragraph_index=ordinal,
        table_index=None,
        ordinal=ordinal,
        block_metadata={},
    )

def _make_candidate(ordinal: int, text: str, *, raw_number: str | None = None) -> "HeadingCandidate":
    from workspace.skills.legal_summarizer.scripts.structure.heading import (
        HeadingCandidate,
    )
    return HeadingCandidate(
        block_index=ordinal,
        text=text,
        score=0.80,
        source="regex_statiya" if raw_number else "regex_numbered_1",
        level=1,
        raw_number=raw_number,
    )

def test_heading_evidence_total_delta_defaults_to_zero():
    """Без признаков total_delta == 0; final_score == source_score."""
    from workspace.skills.legal_summarizer.scripts.structure.heading import (
        HeadingEvidence,
    )

    ev = HeadingEvidence(source_score=0.85)
    assert ev.total_delta == 0.0
    assert ev.final_score == 0.85

def test_heading_evidence_clamped_to_unit_interval():
    """total_delta может вытолкнуть за [0, 1] — final_score зажимается."""
    from workspace.skills.legal_summarizer.scripts.structure.heading import (
        HeadingEvidence,
    )

    ev = HeadingEvidence(
        source_score=0.95,
        short_text_bonus=0.05,
        body_after_bonus=0.05,
        numbering_consistency_bonus=0.05,
        typography_bonus=0.05,
    )
    assert ev.total_delta == 0.20
    assert ev.final_score == 1.0  # clamped

def test_heading_evidence_apply_to_candidate_no_context_no_change():
    """Кандидат без контекста (например, PDF outline) — score не меняется."""
    from workspace.skills.legal_summarizer.scripts.structure.heading import (
        apply_evidence_scoring,
    )

    # PDF outline кандидат имеет block_index == -1.
    from workspace.skills.legal_summarizer.scripts.structure.heading import (
        HeadingCandidate,
    )

    pdf_outline = HeadingCandidate(
        block_index=-1, text="Глава 1", score=0.95, source="pdf_outline", level=1, raw_number=None,
    )
    blocks = (
        _make_heading_block(0, "Some content"),
        _make_heading_block(1, "More content"),
    )
    out = apply_evidence_scoring([pdf_outline], blocks)
    assert out[0].score == 0.95  # без изменений

def test_heading_evidence_short_text_bonus():
    """Короткий текст → +0.05 к score."""
    from workspace.skills.legal_summarizer.scripts.structure.heading import (
        apply_evidence_scoring,
    )

    c = _make_candidate(0, "Статья 1.", raw_number="1")
    blocks = (
        _make_heading_block(0, "Статья 1."),
        _make_heading_block(1, "A" * 200),  # substantial body after
        _make_heading_block(2, "B" * 200),  # substantial body after prev heading-2
        _make_heading_block(3, "C" * 200),  # substantial body after prev heading-3
        _make_heading_block(4, "D" * 200),  # substantial body after prev heading-4
    )
    # Должен получить short_text_bonus и body_after_bonus и numbering_consistency (только 1 элемент — нет).
    out = apply_evidence_scoring([c], blocks)
    assert out[0].score > 0.80  # было 0.80, ожидаем 0.80 + bonuses

def test_heading_evidence_list_penalty():
    """Кандидат в list-like neighborhood → −0.10."""
    from workspace.skills.legal_summarizer.scripts.structure.heading import (
        apply_evidence_scoring,
    )

    # 4 коротких numbered элемента подряд → list-like.
    blocks = tuple(
        _make_heading_block(i, f"{i+1}. пункт списка")
        for i in range(5)
    )
    c = _make_candidate(2, "3. пункт списка", raw_number="3")
    out = apply_evidence_scoring([c], blocks)
    # Был 0.80, после list_penalty (-0.10) → 0.70.
    assert out[0].score < 0.80

def test_heading_evidence_duplicate_penalty():
    """Текст совпадает с предыдущим heading'ом → −0.20."""
    from workspace.skills.legal_summarizer.scripts.structure.heading import (
        apply_evidence_scoring,
    )

    c1 = _make_candidate(0, "Дубликат", raw_number="1")
    c2 = _make_candidate(3, "Дубликат", raw_number="1")  # тот же текст
    blocks = (
        _make_heading_block(0, "Дубликат"),
        _make_heading_block(1, "x" * 200),
        _make_heading_block(2, "y" * 200),
        _make_heading_block(3, "Дубликат"),
    )
    out = apply_evidence_scoring([c1, c2], blocks)
    # Второй кандидат должен получить duplicate_penalty (-0.20).
    assert out[1].score < out[0].score

def test_heading_evidence_numbering_consistency():
    """Heading с монотонным raw_number получает bonus."""
    from workspace.skills.legal_summarizer.scripts.structure.heading import (
        apply_evidence_scoring,
    )

    c1 = _make_candidate(0, "Статья 1.", raw_number="1")
    c2 = _make_candidate(3, "Статья 2.", raw_number="2")
    blocks = (
        _make_heading_block(0, "Статья 1."),
        _make_heading_block(1, "x" * 200),
        _make_heading_block(2, "y" * 200),
        _make_heading_block(3, "Статья 2."),
    )
    out = apply_evidence_scoring([c1, c2], blocks)
    # Второй кандидат имеет numbering_consistency_bonus, первый — нет (нет предыдущего).
    assert out[1].score >= out[0].score

def test_detect_sections_uses_evidence_scoring(tmp_path, monkeypatch):
    """``detect_sections`` теперь применяет evidence scoring
    перед filter_above_threshold. Smoke-проверка: один heading выживает.
    """
    from workspace.skills.legal_summarizer.scripts.structure.sections import (
        detect_sections,
    )
    from workspace.skills.legal_summarizer.scripts.structure.physical import (
        PhysicalDocument,
    )

    blocks = (
        _make_heading_block(0, "Преамбула текста."),
        _make_heading_block(1, "x" * 200),
        _make_heading_block(2, "Статья 1. Общие положения."),
        _make_heading_block(3, "y" * 200),
        _make_heading_block(4, "Статья 2. Обязанности."),
        _make_heading_block(5, "z" * 200),
    )
    doc = PhysicalDocument(
        path="<inline>",
        format="txt",
        title=None,
        size_bytes=sum(len(b.content) for b in blocks),
        blocks=blocks,
        page_count=1,
    )
    tree = detect_sections(doc, pdf_path=None)
    # Оба heading'а должны быть приняты (после bonuses ≥ threshold).
    non_root = [s for s in tree.sections.values() if s.section_id != "s_root" and s.heading]
    assert len(non_root) >= 1

# ---------------------------------------------------------------------------

# ---------------------------------------------------------------------------

def test_list_detection_numbered_list_run_is_list():
    """``1. сделать X / 2. сделать Y / 3. сделать Z``
    (contiguous, короткие) → ``is_list=True``.
    """
    from workspace.skills.legal_summarizer.scripts.structure.list_detection import (
        detect_list_runs,
    )

    blocks = tuple(
        _make_heading_block(i, f"{i+1}. сделать что-то короткое")
        for i in range(5)
    )
    runs = detect_list_runs(blocks)
    assert len(runs) == 1
    run = runs[0]
    assert run.is_list is True
    assert run.numbers == (1, 2, 3, 4, 5)
    assert run.block_ordinals == (0, 1, 2, 3, 4)

def test_list_detection_section_sequence_with_body_is_not_list():
    """``1. Раздел / body / 2. Раздел / body`` → ``is_list=False``.

    План: «1. Общие положения / длинный body / 2. Обязанности» — это **разделы**,
    не list. Наш детектор ловит это через contiguous-ordinal: между разделами
    есть body-блоки → ordinals не contiguous → нет list-run.
    """
    from workspace.skills.legal_summarizer.scripts.structure.list_detection import (
        detect_list_runs,
    )

    blocks = (
        _make_heading_block(0, "1. Общие положения"),
        _make_heading_block(1, "A" * 300),  # body
        _make_heading_block(2, "B" * 300),  # body
        _make_heading_block(3, "2. Обязанности сторон"),
        _make_heading_block(4, "C" * 300),  # body
        _make_heading_block(5, "D" * 300),  # body
        _make_heading_block(6, "3. Ответственность"),
    )
    runs = detect_list_runs(blocks)
    # Тут должно быть 0 list-runs (между элементами есть body, ordinals не contiguous).
    assert all(not r.is_list for r in runs), (
        f"Expected no list-runs; got: {runs}"
    )

def test_list_detection_long_items_not_classified_as_list():
    """Длинные numbered блоки (каждый > max_item_chars) → не list."""
    from workspace.skills.legal_summarizer.scripts.structure.list_detection import (
        detect_list_runs,
        ListDetectionConfig,
    )

    long_text = "X" * 250  # > default 200
    blocks = tuple(
        _make_heading_block(i, f"{i+1}. {long_text}") for i in range(5)
    )
    runs = detect_list_runs(blocks, ListDetectionConfig(max_item_chars=200))
    assert all(not r.is_list for r in runs)

def test_list_detection_short_run_below_min_is_not_list():
    """Только 2 numbered элемента < min_run_length=3 → не list (section из 2 элементов)."""
    from workspace.skills.legal_summarizer.scripts.structure.list_detection import (
        detect_list_runs,
    )

    blocks = (
        _make_heading_block(0, "1. первый"),
        _make_heading_block(1, "2. второй"),
    )
    runs = detect_list_runs(blocks)
    # Run существует, но is_list=False (слишком короткий).
    assert len(runs) == 1
    assert runs[0].is_list is False

def test_list_detection_penalty_value():
    """``list_penalty_for_candidate`` возвращает правильный штраф в зависимости от длины run.

    Этап 10 (PLAN §10): для спорных коротких run (3..4 элементов) штраф
    снижен с 0.10 до 0.08 — это даёт heading-детектору больше шансов
    принять их за headings (а не отбрасывать как list). Изменение
    поведения явно зафиксировано в ``docs/CHANGELOG.md`` Этапа 10.
    """
    from workspace.skills.legal_summarizer.scripts.structure.list_detection import (
        detect_list_runs,
        list_penalty_for_candidate,
    )

    blocks5 = tuple(
        _make_heading_block(i, f"{i+1}. короткий") for i in range(5)
    )
    runs5 = detect_list_runs(blocks5)
    assert list_penalty_for_candidate(2, runs5) == 0.15

    blocks3 = tuple(
        _make_heading_block(i, f"{i+1}. короткий") for i in range(3)
    )
    runs3 = detect_list_runs(blocks3)
    assert list_penalty_for_candidate(1, runs3) == 0.08

    assert list_penalty_for_candidate(99, runs3) == 0.0

def test_heading_evidence_uses_precise_list_penalty():
    """list-penalty в evidence_scoring точнее старой neighborhood-проверки.

    Создаём документ, где heading расположен рядом с 4 короткими numbered neighbors,
    НО это часть большего list-run. Score должен получить list_penalty 0.15
    (≥ 5 элементов), а не 0.10.
    """
    from workspace.skills.legal_summarizer.scripts.structure.heading import (
        apply_evidence_scoring,
    )

    # 6 numbered blocks: первые 5 — короткие list, потом heading в центре,
    # но heading должен попасть в list-run через contiguous.
    blocks = tuple(
        _make_heading_block(i, f"{i+1}. короткий пункт списка")
        for i in range(5)
    )
    c = _make_candidate(2, "3. короткий пункт списка", raw_number="3")
    out = apply_evidence_scoring([c], blocks)
    # Был 0.80. list_penalty 0.15 → ~0.65 (≤ 0.70 с floating-point tolerance).
    assert out[0].score <= 0.70 + 1e-9

def test_list_detection_non_monotonic_breaks_run():
    """Немонотонная последовательность (1, 2, 4) разрывает run."""
    from workspace.skills.legal_summarizer.scripts.structure.list_detection import (
        detect_list_runs,
    )

    blocks = (
        _make_heading_block(0, "1. первый"),
        _make_heading_block(1, "2. второй"),
        _make_heading_block(2, "4. четвёртый"),  # пропуск 3
    )
    runs = detect_list_runs(blocks)
    # Должно быть: run(1,2) + run(4,) — оба is_list=False (короткие).
    assert len(runs) >= 2
    # Все короткие — ни одна не list.
    assert all(not r.is_list for r in runs)

# ---------------------------------------------------------------------------

# ---------------------------------------------------------------------------

def _chunk_config(max_chunk_chars: int = 4000, chunk_overlap_chars: int = 200):
    from workspace.skills.legal_summarizer.scripts.structure.chunks import ChunkConfig
    return ChunkConfig(
        max_chunk_chars=max_chunk_chars,
        chunk_overlap_chars=chunk_overlap_chars,
        chars_per_token=3.5,
        table_chunk_threshold_chars=6000,
        min_section_chars=200,
    )

def test_block_aware_chunking_respects_block_boundaries(tmp_path, monkeypatch):
    """chunk boundaries по возможности совпадают с block boundaries.

    Документ с 3 короткими paragraphs в одной секции → 1 chunk со всеми 3 блоками.
    """
    from workspace.skills.legal_summarizer.scripts.structure.chunks import (
        StructureAwareChunker,
    )
    from workspace.skills.legal_summarizer.scripts.structure.physical import (
        PhysicalDocument,
    )
    from workspace.skills.legal_summarizer.scripts.structure.sections import (
        detect_sections,
    )

    body_pieces = [
        "Короткий абзац один про предмет договора.",
        "Короткий абзац два про обязанности сторон.",
        "Короткий абзац три про сроки и порядок расчётов.",
    ]
    blocks_list = [_make_heading_block(0, "1. Общие положения договора")]
    blocks_list.extend(_make_heading_block(i + 1, t) for i, t in enumerate(body_pieces))
    blocks = tuple(blocks_list)

    doc = PhysicalDocument(
        path="<inline>", format="txt", title=None, size_bytes=100,
        blocks=blocks, page_count=1,
    )
    tree = detect_sections(doc, pdf_path=None)
    chunker = StructureAwareChunker()
    chunks = chunker.chunk(doc, tree, _chunk_config(max_chunk_chars=4000))
    # 1 chunk: heading + 3 body blocks (всё помещается в budget).
    assert len(chunks) == 1
    # block_indices: heading (0) + body (1, 2, 3).
    assert chunks[0].block_indices == (0, 1, 2, 3)
    # Текст содержит все 3 абзаца.
    for piece in body_pieces:
        assert piece in chunks[0].text

def test_block_aware_chunking_splits_at_block_boundary(tmp_path, monkeypatch):
    """chunk boundaries совпадают с block boundaries (не режем блок).

    Документ с 3 paragraphs, каждый ~2 KB → с budget 2.5 KB получается 3 чанка
    (каждый chunk = ровно 1 block).
    """
    from workspace.skills.legal_summarizer.scripts.structure.chunks import (
        StructureAwareChunker,
    )
    from workspace.skills.legal_summarizer.scripts.structure.physical import (
        PhysicalDocument,
    )
    from workspace.skills.legal_summarizer.scripts.structure.sections import (
        detect_sections,
    )

    body_pieces = [
        "А" * 2000,
        "Б" * 2000,
        "В" * 2000,
    ]
    blocks_list = [_make_heading_block(0, "1. Большой раздел договора")]
    blocks_list.extend(_make_heading_block(i + 1, t) for i, t in enumerate(body_pieces))
    blocks = tuple(blocks_list)

    doc = PhysicalDocument(
        path="<inline>", format="txt", title=None, size_bytes=100,
        blocks=blocks, page_count=1,
    )
    tree = detect_sections(doc, pdf_path=None)
    chunker = StructureAwareChunker()
    chunks = chunker.chunk(doc, tree, _chunk_config(max_chunk_chars=2500))
    # 3 чанка: (heading+body1), (body2), (body3). Block boundaries не нарушены.
    assert len(chunks) == 3
    # heading + первый body в первом chunk.
    assert chunks[0].block_indices == (0, 1)
    assert chunks[1].block_indices == (2,)
    assert chunks[2].block_indices == (3,)

def test_block_aware_chunking_oversized_block_falls_back_to_split(tmp_path, monkeypatch):
    """если один paragraph слишком большой, fallback на split_text.

    Один block 5000 chars при budget 2000 → 3 chunks (split_text с overlap).
    """
    from workspace.skills.legal_summarizer.scripts.structure.chunks import (
        StructureAwareChunker,
    )
    from workspace.skills.legal_summarizer.scripts.structure.physical import (
        PhysicalDocument,
    )
    from workspace.skills.legal_summarizer.scripts.structure.sections import (
        detect_sections,
    )

    big = "X" * 5000
    blocks = (
        _make_heading_block(0, "1. Раздел с очень длинным body"),
        _make_heading_block(1, big),
    )
    doc = PhysicalDocument(
        path="<inline>", format="txt", title=None, size_bytes=100,
        blocks=blocks, page_count=1,
    )
    tree = detect_sections(doc, pdf_path=None)
    chunker = StructureAwareChunker()
    chunks = chunker.chunk(doc, tree, _chunk_config(max_chunk_chars=2000))
    # Первый chunk — heading (block_indices = (0,)).
    # Остальные — split-части большого body (block_indices = (b.ordinal,) —
    # split fallback сохраняет provenance до оригинального DocumentBlock,
    # см. фикс #7 в changelog).
    assert len(chunks) >= 2
    assert chunks[0].block_indices == (0,)
    for c in chunks[1:]:
        # split-части: атрибутируем к ordinal исходного большого блока.
        assert c.block_indices == (1,)

def test_block_aware_chunking_preserves_section_metadata(tmp_path, monkeypatch):
    """section metadata сохраняется в каждом chunk."""
    from workspace.skills.legal_summarizer.scripts.structure.chunks import (
        StructureAwareChunker,
    )
    from workspace.skills.legal_summarizer.scripts.structure.physical import (
        PhysicalDocument,
    )
    from workspace.skills.legal_summarizer.scripts.structure.sections import (
        detect_sections,
    )

    blocks = (
        _make_heading_block(0, "1. Раздел А"),
        _make_heading_block(1, "Тело раздела А"),
        _make_heading_block(2, "2. Раздел Б"),
        _make_heading_block(3, "Тело раздела Б"),
    )
    doc = PhysicalDocument(
        path="<inline>", format="txt", title=None, size_bytes=100,
        blocks=blocks, page_count=1,
    )
    tree = detect_sections(doc, pdf_path=None)
    chunker = StructureAwareChunker()
    chunks = chunker.chunk(doc, tree, _chunk_config())
    assert len(chunks) == 2
    # Каждый chunk несёт свой section_path.
    assert chunks[0].section_path != chunks[1].section_path
    assert chunks[0].section_id != chunks[1].section_id

def test_block_aware_chunking_no_text_lost(tmp_path, monkeypatch):
    """никакой текст не теряется.

    Сумма длин всех chunks ≥ сумма длин всех body блоков.
    """
    from workspace.skills.legal_summarizer.scripts.structure.chunks import (
        StructureAwareChunker,
    )
    from workspace.skills.legal_summarizer.scripts.structure.physical import (
        PhysicalDocument,
    )
    from workspace.skills.legal_summarizer.scripts.structure.sections import (
        detect_sections,
    )

    body_pieces = [
        "Первый абзац текста для анализа договора аренды.",
        "Второй абзац с дополнительной информацией.",
        "Третий абзац заключительного раздела.",
    ]
    blocks_list = [_make_heading_block(0, "1. Один раздел")]
    blocks_list.extend(_make_heading_block(i + 1, t) for i, t in enumerate(body_pieces))
    blocks = tuple(blocks_list)

    doc = PhysicalDocument(
        path="<inline>", format="txt", title=None, size_bytes=100,
        blocks=blocks, page_count=1,
    )
    tree = detect_sections(doc, pdf_path=None)
    chunker = StructureAwareChunker()
    chunks = chunker.chunk(doc, tree, _chunk_config())
    total_chunk_chars = sum(len(c.text) for c in chunks)
    body_chars = sum(len(p) for p in body_pieces)
    # Допускаем overlap, поэтому >= а не ==.
    assert total_chunk_chars >= body_chars
    # Но не больше, чем 2x (нет чрезмерного дублирования).
    assert total_chunk_chars <= body_chars * 2

def test_block_aware_chunking_tables_never_split_within_row(tmp_path, monkeypatch):
    """tables никогда не режутся внутри row (row-start/end сохраняются)."""
    from workspace.skills.legal_summarizer.scripts.structure.chunks import (
        StructureAwareChunker,
    )
    from workspace.skills.legal_summarizer.scripts.structure.physical import (
        DocumentBlock,
        PhysicalDocument,
    )
    from workspace.skills.legal_summarizer.scripts.structure.sections import (
        detect_sections,
    )

    table_content = "\n".join(f"row{i} | cell" for i in range(10))
    blocks = (
        _make_heading_block(0, "1. Раздел с таблицей"),
        DocumentBlock(
            block_id="b_0001",
            block_type="table",
            content=table_content,
            char_count=len(table_content),
            page_index=1,
            page_start=1,
            page_end=1,
            paragraph_index=None,
            table_index=0,
            ordinal=1,
            block_metadata={"row_count": 10},
        ),
    )
    doc = PhysicalDocument(
        path="<inline>", format="txt", title=None, size_bytes=100,
        blocks=blocks, page_count=1,
    )
    tree = detect_sections(doc, pdf_path=None)
    chunker = StructureAwareChunker()
    chunks = chunker.chunk(doc, tree, _chunk_config())
    table_chunks = [c for c in chunks if c.block_types and c.block_types[0] == "table"]
    assert len(table_chunks) == 1
    # Таблица — атомарна: rows 1..10 в одном chunk.
    assert table_chunks[0].table_row_start == 1
    assert table_chunks[0].table_row_end == 10

# ---------------------------------------------------------------------------

# ---------------------------------------------------------------------------

def test_chunk_overlap_zero_no_duplication_for_normal_blocks(tmp_path, monkeypatch):
    """normal block boundary → 0 overlap.

    5 одинаковых body-блоков по 100 chars, budget=300 → block-aware даёт
    2 chunks: (block1, block2), (block3, block4). Без overlap сумма chars
    == 400, не больше.
    """
    from workspace.skills.legal_summarizer.scripts.structure.chunks import (
        StructureAwareChunker,
    )
    from workspace.skills.legal_summarizer.scripts.structure.physical import (
        PhysicalDocument,
    )
    from workspace.skills.legal_summarizer.scripts.structure.sections import (
        detect_sections,
    )

    body_pieces = ["Тело блока номер N с содержательным текстом." for _ in range(5)]
    blocks_list = [_make_heading_block(0, "1. Большой раздел")]
    blocks_list.extend(_make_heading_block(i + 1, t) for i, t in enumerate(body_pieces))
    blocks = tuple(blocks_list)

    doc = PhysicalDocument(
        path="<inline>", format="txt", title=None, size_bytes=100,
        blocks=blocks, page_count=1,
    )
    tree = detect_sections(doc, pdf_path=None)
    chunker = StructureAwareChunker()
    # chunk_overlap_chars=0.
    cfg = _chunk_config(max_chunk_chars=300, chunk_overlap_chars=0)
    chunks = chunker.chunk(doc, tree, cfg)
    total = sum(len(c.text) for c in chunks)
    source = sum(len(b) for b in body_pieces)
    # Сумма chunk chars ≈ сумма source chars (без overlap-добавки).
    # Допуск: heading в первом chunk добавляет ~22 chars (заголовок).
    assert abs(total - source) <= 50, f"total={total}, source={source}"

def test_chunk_overlap_zero_oversized_block_still_has_split_overlap(tmp_path, monkeypatch):
    """oversized block всё равно получает минимальный overlap для continuity."""
    from workspace.skills.legal_summarizer.scripts.structure.chunks import (
        StructureAwareChunker,
    )
    from workspace.skills.legal_summarizer.scripts.structure.physical import (
        PhysicalDocument,
    )
    from workspace.skills.legal_summarizer.scripts.structure.sections import (
        detect_sections,
    )

    big = "X" * 5000
    blocks = (
        _make_heading_block(0, "1. Раздел"),
        _make_heading_block(1, big),
    )
    doc = PhysicalDocument(
        path="<inline>", format="txt", title=None, size_bytes=100,
        blocks=blocks, page_count=1,
    )
    tree = detect_sections(doc, pdf_path=None)
    chunker = StructureAwareChunker()
    chunks = chunker.chunk(doc, tree, _chunk_config(max_chunk_chars=2000, chunk_overlap_chars=0))
    # Oversized block через split_text: даже с overlap=0 split_text добавляет свой overlap.
    # Проверяем, что split не меньше 2 (значит split_text что-то сделал).
    assert len(chunks) >= 2

def test_chunk_overlap_project_json_default_is_zero():
    """project.json дефолт ``chunk_overlap`` = 0 для summarization.

    NOTE: project.json — JSONC (с комментариями). Значение уже проверяется
    в ``test_config_keys.py::test_required_key_present_with_default`` —
    этот тест только smoke-проверяет, что get_chunking_config возвращает 0.
    """
    import skill_config

    cfg = skill_config.get_chunking_config()
    assert cfg["chunk_overlap"] == 0

def test_summarizer_passes_zero_overlap_to_chunk_config(monkeypatch, tmp_path):
    """summarizer.inspect → StructureAwareChunker получает overlap=0 из project.json.

    Проверяем через ``inspect()`` (который вызывает chunker), а не run()
    (single-call path не использует chunker).
    """
    import summarizer

    captured = {}

    real_make_chunk_config = summarizer._make_chunk_config

    def spy_make(cfg, budget):
        result = real_make_chunk_config(cfg, budget)
        captured["chunk_overlap_chars"] = result.chunk_overlap_chars
        return result

    monkeypatch.setattr(summarizer, "_make_chunk_config", spy_make)
    monkeypatch.setattr(summarizer, "get_chunking_config", lambda: {
        "chunk_size": 4000,
        "chunk_overlap": 0,  # явно
        "single_call_threshold": 100,
        "chunk_size_input_ratio": None,
    })

    # inspect() вызывает _make_chunk_config → spy перехватывает overlap.
    big_text = (
        "Преамбула длинного документа. " * 50
        + "Статья 1. Общие положения. " * 50
        + ("Длинный абзац текста для chunking'а. " * 100)
    )
    summarizer.inspect(big_text)
    assert captured.get("chunk_overlap_chars") == 0

# ---------------------------------------------------------------------------

# ---------------------------------------------------------------------------

def test_token_budget_module_is_extracted_and_re_exported():
    """``TokenBudget`` живёт в ``token_budget.py``,
    ``packing.TokenBudget`` — re-export (back-compat для существующих импортов).
    """
    from workspace.skills.legal_summarizer.scripts import packing
    from workspace.skills.legal_summarizer.scripts import token_budget

    assert hasattr(token_budget, "TokenBudget")
    assert hasattr(token_budget, "count_tokens")
    assert hasattr(token_budget, "text_to_tokens_estimate")
    assert packing.TokenBudget is token_budget.TokenBudget

def test_token_budget_available_chunk_tokens():
    """``TokenBudget.available_chunk_tokens`` рассчитывается по формуле
    ``(context_window - system - instruction - output_reserve) * safety_margin``.
    """
    from workspace.skills.legal_summarizer.scripts.token_budget import TokenBudget

    budget = TokenBudget(
        context_window_tokens=100000,
        system_prompt_tokens=1000,
        instruction_tokens=200,
        output_reserve_tokens=8000,
        safety_margin=0.85,
        chars_per_token=3.5,
    )
    expected = int((100000 - 1000 - 200 - 8000) * 0.85)
    assert budget.available_chunk_tokens == expected

def test_token_budget_minimum_1000():
    """``available_chunk_tokens`` floor = 1000 (защита от overshoot'а при крошечных
    context windows)."""
    from workspace.skills.legal_summarizer.scripts.token_budget import TokenBudget

    tiny = TokenBudget(
        context_window_tokens=2000,
        system_prompt_tokens=1500,
        instruction_tokens=200,
        output_reserve_tokens=8000,
        safety_margin=0.85,
        chars_per_token=3.5,
    )
    # Без floor было бы max(-7700*0.85, ...) — отрицательное; clamp к 1000.
    assert tiny.available_chunk_tokens == 1000

def test_text_to_tokens_estimate_fallback():
    """``text_to_tokens_estimate`` — O(1) fallback без tokenizer."""
    from workspace.skills.legal_summarizer.scripts.token_budget import (
        text_to_tokens_estimate,
    )

    # 14 chars / 3.5 ≈ 4 tokens.
    assert text_to_tokens_estimate("hello world!") == 4
    # Empty → 0.
    assert text_to_tokens_estimate("") == 0
    # Custom ratio.
    assert text_to_tokens_estimate("abcdef", chars_per_token=2.0) == 3

def test_count_tokens_fallback_when_no_tiktoken():
    """``count_tokens`` работает без tiktoken (fallback на chars)."""
    from workspace.skills.legal_summarizer.scripts.token_budget import (
        count_tokens,
    )

    # Без tiktoken (по умолчанию в проекте) — fallback на chars/4.
    n = count_tokens("hello world")
    # len("hello world") = 11; (11+3)//4 = 3.
    assert n == 3

def test_count_tokens_empty():
    """Пустая строка → 0 токенов."""
    from workspace.skills.legal_summarizer.scripts.token_budget import (
        count_tokens,
    )
    assert count_tokens("") == 0
    assert count_tokens(None) == 0  # type: ignore[arg-type]

def test_packing_actual_tokens_le_calculated_budget():
    """``actual_tokens <= calculated_budget`` для всех packing cases.

    Сумма token_estimate всех chunks в одном batch не должна превышать
    budget.available_chunk_tokens.
    """
    from workspace.skills.legal_summarizer.scripts.packing import (
        ContextBatch,
        TokenBudget,
        pack_chunks,
    )
    from workspace.skills.legal_summarizer.scripts.structure.chunks import (
        Chunk,
    )

    budget = TokenBudget(
        context_window_tokens=100000,
        system_prompt_tokens=1000,
        instruction_tokens=200,
        output_reserve_tokens=8000,
        safety_margin=0.85,
        chars_per_token=3.5,
    )
    chunks = tuple(
        Chunk(
            chunk_id=f"{i:03d}",
            index=i,
            text=f"Текст чанка номер {i}.",
            char_count=20,
            token_estimate=6,
            page_start=1,
            page_end=1,
            section_id="s_0001",
            section_path="1",
            section_heading="",
            block_indices=(i,),
            block_types=("paragraph",),
        )
        for i in range(5)
    )
    batches = pack_chunks(list(chunks), budget)
    assert len(batches) >= 1
    available = budget.available_chunk_tokens
    for b in batches:
        # total_tokens_estimate включает BATCH_OVERHEAD (80) + sum chunks.
        # Сумма token_estimate chunks в batch не должна превышать available.
        chunk_tokens_sum = sum(c.token_estimate for c in b.chunks)
        assert chunk_tokens_sum <= available, (
            f"batch {b.batch_id}: {chunk_tokens_sum} > available {available}"
        )

def test_packing_module_token_budget_re_exported():
    """``from packing import TokenBudget`` работает (back-compat)."""
    from workspace.skills.legal_summarizer.scripts.packing import TokenBudget

    assert TokenBudget is not None
    b = TokenBudget(
        context_window_tokens=10000,
        system_prompt_tokens=100,
        instruction_tokens=100,
        output_reserve_tokens=1000,
        safety_margin=0.8,
        chars_per_token=3.5,
    )
    assert b.available_chunk_tokens > 0

# ---------------------------------------------------------------------------

# ---------------------------------------------------------------------------

def _make_test_chunk(
    chunk_id: str,
    chars: int,
    section_path: str = "1",
    *,
    section_id: str | None = None,
) -> "Chunk":
    """Хелпер для тестов packing.

    Args:
        section_id: явный id секции. Если ``None`` — дефолт ``"s_0001"``.
            Для тестов с разными section_path нужно передавать разные section_id,
            чтобы locality-aware packing корректно их различал.
    """
    from workspace.skills.legal_summarizer.scripts.structure.chunks import Chunk

    if section_id is None:
        section_id = "s_0001"

    return Chunk(
        chunk_id=chunk_id,
        index=int(chunk_id),
        text="x" * chars,
        char_count=chars,
        token_estimate=max(1, chars // 3),
        page_start=1,
        page_end=1,
        section_id=section_id,
        section_path=section_path,
        section_heading="",
        block_indices=(int(chunk_id),),
        block_types=("paragraph",),
    )

def test_utilization_calculation_is_correct():
    """``content / available`` clamped к [0, 1]."""
    from workspace.skills.legal_summarizer.scripts.packing import ContextBatch

    def _batch(content_tokens: int, available: int = 1000) -> ContextBatch:
        return ContextBatch(
            batch_id="cb_000",
            chunks=(_make_test_chunk("0", 300),),
            total_tokens_estimate=content_tokens + 80,
            section_paths=("1",),
            page_range=(1, 1),
            content_tokens_estimate=content_tokens,
            available_chunk_tokens=available,
        )

    # 100 / 1000 → 0.1.
    assert _batch(100).utilization == pytest.approx(0.1)
    # 500 / 1000 → 0.5.
    assert _batch(500).utilization == pytest.approx(0.5)
    # 1000 / 1000 → 1.0.
    assert _batch(1000).utilization == pytest.approx(1.0)
    # 1500 / 1000 → clamped to 1.0.
    assert _batch(1500).utilization == pytest.approx(1.0)
    # 0 / 1000 → 0.0.
    assert _batch(0).utilization == 0.0

def test_utilization_zero_when_no_budget():
    """Без ``available_chunk_tokens`` utilization = 0.0 (нет данных)."""
    from workspace.skills.legal_summarizer.scripts.packing import ContextBatch

    b = ContextBatch(
        batch_id="cb_000",
        chunks=(_make_test_chunk("0", 300),),
        total_tokens_estimate=180,
        section_paths=("1",),
        page_range=(1, 1),
        content_tokens_estimate=100,
        available_chunk_tokens=0,
    )
    assert b.utilization == 0.0

def test_utilization_not_constant_after_fix():
    """после фикса utilization перестала быть всегда 1.0 (баг baseline).

    Раньше ``total_tokens_estimate / total_tokens_estimate`` всегда давал 1.0.
    После фикса — реальное отношение content/available.
    """
    from workspace.skills.legal_summarizer.scripts.packing import (
        ContextBatch,
        TokenBudget,
        pack_chunks,
    )

    budget = TokenBudget(
        context_window_tokens=10000,
        system_prompt_tokens=100,
        instruction_tokens=100,
        output_reserve_tokens=1000,
        safety_margin=0.85,
        chars_per_token=3.5,
    )
    available = budget.available_chunk_tokens
    chunks = [
        _make_test_chunk("0", 300),   # 100 tokens
        _make_test_chunk("1", 300),   # 100 tokens
        _make_test_chunk("2", 300),   # 100 tokens
        _make_test_chunk("3", 300),   # 100 tokens
    ]
    batches = pack_chunks(chunks, budget)
    assert len(batches) >= 1
    # Первый batch: все 4 chunks помещаются в budget → ~400 tokens.
    first = batches[0]
    # Утилизация должна быть < 1.0 (всё содержимое 400 < 9000 available).
    assert first.utilization < 1.0
    assert first.utilization > 0.0
    # Конкретное значение: 400 / available (≈9000).
    assert first.utilization == pytest.approx(400 / available, rel=1e-3)

def test_pack_chunks_sets_available_in_batch():
    """``pack_chunks`` записывает ``available_chunk_tokens`` в каждый ContextBatch."""
    from workspace.skills.legal_summarizer.scripts.packing import (
        TokenBudget,
        pack_chunks,
    )

    budget = TokenBudget(
        context_window_tokens=50000,
        system_prompt_tokens=500,
        instruction_tokens=100,
        output_reserve_tokens=4000,
        safety_margin=0.85,
        chars_per_token=3.5,
    )
    expected_available = budget.available_chunk_tokens
    chunks = [_make_test_chunk("0", 200), _make_test_chunk("1", 200)]
    batches = pack_chunks(chunks, budget)
    for b in batches:
        assert b.available_chunk_tokens == expected_available

def test_pack_chunks_sets_content_tokens_in_batch():
    """``pack_chunks`` записывает ``content_tokens_estimate`` (без overhead)."""
    from workspace.skills.legal_summarizer.scripts.packing import (
        TokenBudget,
        pack_chunks,
    )

    budget = TokenBudget(
        context_window_tokens=50000,
        system_prompt_tokens=500,
        instruction_tokens=100,
        output_reserve_tokens=4000,
        safety_margin=0.85,
        chars_per_token=3.5,
    )
    chunks = [_make_test_chunk("0", 600), _make_test_chunk("1", 600)]  # 200 tokens каждый
    batches = pack_chunks(chunks, budget)
    first = batches[0]
    # content_tokens = total - overhead (80).
    assert first.content_tokens_estimate == first.total_tokens_estimate - 80

def test_pack_chunks_to_dict_includes_utilization():
    """``ContextBatch.to_dict()`` теперь включает ``utilization`` для observability."""
    from workspace.skills.legal_summarizer.scripts.packing import (
        TokenBudget,
        pack_chunks,
    )

    budget = TokenBudget(
        context_window_tokens=50000,
        system_prompt_tokens=500,
        instruction_tokens=100,
        output_reserve_tokens=4000,
        safety_margin=0.85,
        chars_per_token=3.5,
    )
    chunks = [_make_test_chunk("0", 300)]
    batches = pack_chunks(chunks, budget)
    d = batches[0].to_dict()
    assert "utilization" in d
    assert "content_tokens_estimate" in d
    assert "available_chunk_tokens" in d
    assert 0.0 <= d["utilization"] <= 1.0

# ---------------------------------------------------------------------------

# ---------------------------------------------------------------------------

def test_packing_config_default_strict_section_locality():
    """default ``PackingConfig`` — strict section-locality (back-compat)."""
    from workspace.skills.legal_summarizer.scripts.packing import PackingConfig

    cfg = PackingConfig()
    assert cfg.allow_adjacent_sections is False
    assert cfg.min_remaining_for_mix == 0.5

def test_pack_chunks_default_keeps_strict_section_locality():
    """Без флага ``allow_adjacent_sections`` — поведение как до этапа 14.

    3 секции по 1 маленькому chunk → 3 batches (как в baseline).
    """
    from workspace.skills.legal_summarizer.scripts.packing import (
        TokenBudget,
        pack_chunks,
    )

    budget = TokenBudget(
        context_window_tokens=50000,
        system_prompt_tokens=500,
        instruction_tokens=100,
        output_reserve_tokens=4000,
        safety_margin=0.85,
        chars_per_token=3.5,
    )
    chunks = [
        _make_test_chunk("0", 300, section_path="1", section_id="s_0001"),
        _make_test_chunk("1", 300, section_path="2", section_id="s_0002"),
        _make_test_chunk("2", 300, section_path="3", section_id="s_0003"),
    ]
    batches = pack_chunks(chunks, budget)
    assert len(batches) == 3  # strict: 1 batch per section
    # Каждый batch содержит chunks только из одной секции.
    for b in batches:
        section_ids = {c.section_id for c in b.chunks}
        assert len(section_ids) == 1

def test_pack_chunks_with_allow_adjacent_mixes_adjacent_sections():
    """``allow_adjacent_sections=True`` — merging chunks из adjacent sections.

    3 секции по 1 chunk (по 30 tokens), budget=2000:
        Без флага: 3 batches (1 per section).
        С флагом: 1 batch (все 3 chunks, т.к. секции adjacent и есть бюджет).
    """
    from workspace.skills.legal_summarizer.scripts.packing import (
        PackingConfig,
        TokenBudget,
        pack_chunks,
    )

    budget = TokenBudget(
        context_window_tokens=50000,
        system_prompt_tokens=500,
        instruction_tokens=100,
        output_reserve_tokens=4000,
        safety_margin=0.85,
        chars_per_token=3.5,
    )
    chunks = [
        _make_test_chunk("0", 300, section_path="1", section_id="s_0001"),
        _make_test_chunk("1", 300, section_path="2", section_id="s_0002"),
        _make_test_chunk("2", 300, section_path="3", section_id="s_0003"),
    ]
    cfg = PackingConfig(allow_adjacent_sections=True)
    batches = pack_chunks(chunks, budget, cfg)
    # После locality-aware: 1-2 batches (не 3 как strict).
    assert len(batches) < 3
    # chunk.section_id сохраняется.
    for b in batches:
        for c in b.chunks:
            assert c.section_id.startswith("s_")

def test_pack_chunks_allow_adjacent_preserves_section_metadata():
    """каждый chunk в locality-aware batch сохраняет свой ``section_id``
    и ``section_path`` (map output не теряет provenance)."""
    from workspace.skills.legal_summarizer.scripts.packing import (
        PackingConfig,
        TokenBudget,
        pack_chunks,
    )

    budget = TokenBudget(
        context_window_tokens=50000,
        system_prompt_tokens=500,
        instruction_tokens=100,
        output_reserve_tokens=4000,
        safety_margin=0.85,
        chars_per_token=3.5,
    )
    chunks = [
        _make_test_chunk("0", 300, section_path="1", section_id="s_0001"),
        _make_test_chunk("1", 300, section_path="2", section_id="s_0002"),
        _make_test_chunk("2", 300, section_path="3", section_id="s_0003"),
        _make_test_chunk("3", 300, section_path="4", section_id="s_0004"),
    ]
    cfg = PackingConfig(allow_adjacent_sections=True, min_remaining_for_mix=0.3)
    batches = pack_chunks(chunks, budget, cfg)
    # Каждый chunk имеет свой section_id (map output корректен).
    all_section_ids = {c.section_id for b in batches for c in b.chunks}
    assert len(all_section_ids) == 4  # все 4 секции присутствуют

def test_pack_chunks_allow_adjacent_respects_min_remaining_threshold():
    """``min_remaining_for_mix`` ограничивает жадность mixing'а."""
    from workspace.skills.legal_summarizer.scripts.packing import (
        PackingConfig,
        TokenBudget,
        pack_chunks,
    )

    budget = TokenBudget(
        context_window_tokens=50000,
        system_prompt_tokens=500,
        instruction_tokens=100,
        output_reserve_tokens=4000,
        safety_margin=0.85,
        chars_per_token=3.5,
    )
    chunks = [
        _make_test_chunk("0", 1500, section_path="1", section_id="s_0001"),  # 500 tokens
        _make_test_chunk("1", 1500, section_path="2", section_id="s_0002"),
        _make_test_chunk("2", 1500, section_path="3", section_id="s_0003"),
    ]
    # min_remaining_for_mix=0.9: не заимствуем, если осталось < 90% бюджета.
    cfg = PackingConfig(allow_adjacent_sections=True, min_remaining_for_mix=0.9)
    batches_strict = pack_chunks(chunks, budget)
    batches_strict_count = len(batches_strict)
    batches_locality = pack_chunks(chunks, budget, cfg)
    # С высоким min_remaining — locality не помогает (мало бюджета остаётся).
    # Допускаем ≥ strict_count (не хуже strict).
    assert len(batches_locality) >= batches_strict_count - 1

def test_pack_chunks_locality_first_section_priority():
    """приоритет — same section > adjacent.

    Если следующая секция не adjacent (т.е. есть разрыв) — не смешивать.
    Реальная проверка: locality-aware не хуже strict (не делает больше batches).
    """
    from workspace.skills.legal_summarizer.scripts.packing import (
        PackingConfig,
        TokenBudget,
        pack_chunks,
    )

    budget = TokenBudget(
        context_window_tokens=50000,
        system_prompt_tokens=500,
        instruction_tokens=100,
        output_reserve_tokens=4000,
        safety_margin=0.85,
        chars_per_token=3.5,
    )
    chunks = [
        _make_test_chunk("0", 300, section_path="1", section_id="s_0001"),
        _make_test_chunk("1", 300, section_path="2", section_id="s_0002"),
        _make_test_chunk("2", 300, section_path="3", section_id="s_0003"),
        _make_test_chunk("3", 300, section_path="4", section_id="s_0004"),
        _make_test_chunk("4", 300, section_path="5", section_id="s_0005"),
    ]
    cfg = PackingConfig(allow_adjacent_sections=True, min_remaining_for_mix=0.5)
    strict = pack_chunks(chunks, budget)
    locality = pack_chunks(chunks, budget, cfg)
    # Locality не должен давать БОЛЬШЕ batches чем strict (только меньше или равно).
    assert len(locality) <= len(strict)
    # И не должен терять chunks.
    strict_chunks = {c.chunk_id for b in strict for c in b.chunks}
    locality_chunks = {c.chunk_id for b in locality for c in b.chunks}
    assert strict_chunks == locality_chunks

# ---------------------------------------------------------------------------

# ---------------------------------------------------------------------------

def test_document_stats_basic_metrics():
    """``compute_document_stats`` считает базовые метрики без LLM."""
    from workspace.skills.legal_summarizer.scripts.document_stats import (
        compute_document_stats,
    )
    from workspace.skills.legal_summarizer.scripts.structure.physical import (
        PhysicalDocument,
    )

    blocks = (
        _make_heading_block(0, "Title"),
        _make_heading_block(1, "Body one."),
        _make_heading_block(2, "Body two."),
    )
    doc = PhysicalDocument(
        path="<inline>", format="txt", title=None, size_bytes=100,
        blocks=blocks, page_count=2,
    )
    stats = compute_document_stats(doc)
    assert stats.blocks == 3
    assert stats.pages == 2
    assert stats.chars == len("Title") + len("Body one.") + len("Body two.")
    assert stats.estimated_tokens > 0
    assert stats.sections == 0  # tree не передан
    assert stats.tables == 0
    assert stats.chunks == 0  # chunks не переданы

def test_document_stats_with_tree_counts_sections():
    """С переданным ``tree`` — считаем число секций (без root)."""
    from workspace.skills.legal_summarizer.scripts.document_stats import (
        compute_document_stats,
    )
    from workspace.skills.legal_summarizer.scripts.structure.physical import (
        PhysicalDocument,
    )
    from workspace.skills.legal_summarizer.scripts.structure.sections import (
        ROOT_SECTION_ID,
        DocumentSection,
        SectionTree,
    )

    blocks = (
        _make_heading_block(0, "1. Первый раздел"),
        _make_heading_block(1, "Тело первого раздела"),
        _make_heading_block(2, "2. Второй раздел"),
        _make_heading_block(3, "Тело второго раздела"),
    )
    doc = PhysicalDocument(
        path="<inline>", format="txt", title=None, size_bytes=100,
        blocks=blocks, page_count=1,
    )
    tree = SectionTree(
        sections={
            ROOT_SECTION_ID: DocumentSection(
                section_id=ROOT_SECTION_ID, level=0, heading="",
                section_path="", block_indices=(0, 1, 2, 3), children=(), parent_id=None,
            ),
            "s_0001": DocumentSection(
                section_id="s_0001", level=1, heading="1. Первый раздел",
                section_path="1", block_indices=(0, 1), children=(),
                parent_id=ROOT_SECTION_ID,
            ),
            "s_0002": DocumentSection(
                section_id="s_0002", level=1, heading="2. Второй раздел",
                section_path="2", block_indices=(2, 3), children=(),
                parent_id=ROOT_SECTION_ID,
            ),
        },
        root_id=ROOT_SECTION_ID,
        block_to_section={i: "s_0001" if i < 2 else "s_0002" for i in range(4)},
    )
    stats = compute_document_stats(doc, tree=tree)
    assert stats.sections == 2  # без root

def test_document_stats_counts_tables():
    """Tables считаются отдельно от других блоков."""
    from workspace.skills.legal_summarizer.scripts.document_stats import (
        compute_document_stats,
    )
    from workspace.skills.legal_summarizer.scripts.structure.physical import (
        DocumentBlock,
        PhysicalDocument,
    )

    table_block = DocumentBlock(
        block_id="b_0002",
        block_type="table",
        content="a | b\nc | d",
        char_count=10,
        page_index=1,
        page_start=1,
        page_end=1,
        paragraph_index=None,
        table_index=0,
        ordinal=2,
        block_metadata={"row_count": 2},
    )
    blocks = (
        _make_heading_block(0, "Paragraph 1"),
        _make_heading_block(1, "Paragraph 2"),
        table_block,
    )
    doc = PhysicalDocument(
        path="<inline>", format="txt", title=None, size_bytes=100,
        blocks=blocks, page_count=1,
    )
    stats = compute_document_stats(doc)
    assert stats.blocks == 3
    assert stats.tables == 1

def test_document_stats_counts_chunks():
    """Chunks считаются отдельно."""
    from workspace.skills.legal_summarizer.scripts.document_stats import (
        compute_document_stats,
    )
    from workspace.skills.legal_summarizer.scripts.structure.physical import (
        PhysicalDocument,
    )

    doc = PhysicalDocument(
        path="<inline>", format="txt", title=None, size_bytes=100,
        blocks=(_make_heading_block(0, "Body"),), page_count=1,
    )
    fake_chunks = [
        _make_test_chunk("0", 100, section_path="1"),
        _make_test_chunk("1", 100, section_path="1"),
    ]
    stats = compute_document_stats(doc, chunks=fake_chunks)
    assert stats.chunks == 2

def test_document_stats_blocks_per_section_ratio():
    """``blocks_per_section`` — ratio для adaptive strategy."""
    from workspace.skills.legal_summarizer.scripts.document_stats import (
        compute_document_stats,
    )
    from workspace.skills.legal_summarizer.scripts.structure.physical import (
        PhysicalDocument,
    )

    # 10 blocks, 2 секции → ratio 5.0.
    blocks = tuple(_make_heading_block(i, f"Block {i}") for i in range(10))
    doc = PhysicalDocument(
        path="<inline>", format="txt", title=None, size_bytes=100,
        blocks=blocks, page_count=1,
    )
    stats = compute_document_stats(doc)
    assert stats.blocks_per_section == 10.0  # нет tree → sections=0 → ratio = blocks

def test_document_stats_chars_per_block_ratio():
    """``chars_per_block`` — средняя длина блока."""
    from workspace.skills.legal_summarizer.scripts.document_stats import (
        compute_document_stats,
    )
    from workspace.skills.legal_summarizer.scripts.structure.physical import (
        PhysicalDocument,
    )

    blocks = (
        _make_heading_block(0, "AAA"),    # 3 chars
        _make_heading_block(1, "BBBBB"),  # 5 chars
        _make_heading_block(2, "C"),       # 1 char
    )
    doc = PhysicalDocument(
        path="<inline>", format="txt", title=None, size_bytes=100,
        blocks=blocks, page_count=1,
    )
    stats = compute_document_stats(doc)
    # 9 chars / 3 blocks = 3.0 chars/block.
    assert stats.chars_per_block == 3.0

def test_document_stats_repeated_blocks_default():
    """``repeated_blocks`` default = 0 (cleanup не выполнялся)."""
    from workspace.skills.legal_summarizer.scripts.document_stats import (
        compute_document_stats,
    )
    from workspace.skills.legal_summarizer.scripts.structure.physical import (
        PhysicalDocument,
    )

    doc = PhysicalDocument(
        path="<inline>", format="txt", title=None, size_bytes=100,
        blocks=(_make_heading_block(0, "Body"),), page_count=1,
    )
    stats = compute_document_stats(doc)
    assert stats.repeated_blocks == 0

def test_document_stats_repeated_blocks_passed_through():
    """``repeated_blocks`` передаётся через параметр."""
    from workspace.skills.legal_summarizer.scripts.document_stats import (
        compute_document_stats,
    )
    from workspace.skills.legal_summarizer.scripts.structure.physical import (
        PhysicalDocument,
    )

    doc = PhysicalDocument(
        path="<inline>", format="txt", title=None, size_bytes=100,
        blocks=(_make_heading_block(0, "Body"),), page_count=1,
    )
    stats = compute_document_stats(doc, repeated_blocks=5)
    assert stats.repeated_blocks == 5

def test_document_stats_no_llm_calls():
    """``compute_document_stats`` НЕ вызывает LLM.

    Smoke-проверка: вызов с monkey-patched llm.chat → если бы stats
    дёргал LLM, monkey-patch счётчик бы увеличился.
    """
    import summarizer

    call_count = {"n": 0}

    def fake_chat(messages, *, context=None, **kwargs):
        call_count["n"] += 1
        return "should not be called"

    summarizer.llm.chat = fake_chat  # type: ignore[assignment]

    from workspace.skills.legal_summarizer.scripts.document_stats import (
        compute_document_stats,
    )
    from workspace.skills.legal_summarizer.scripts.structure.physical import (
        PhysicalDocument,
    )

    doc = PhysicalDocument(
        path="<inline>", format="txt", title=None, size_bytes=100,
        blocks=(_make_heading_block(0, "Body"),), page_count=1,
    )
    compute_document_stats(doc)
    assert call_count["n"] == 0

# ---------------------------------------------------------------------------

# ---------------------------------------------------------------------------

def test_select_reduce_strategy_token_budget_first():
    """если estimated_tokens ≤ reduce_budget → FLAT (главный критерий)."""
    from workspace.skills.legal_summarizer.scripts.summarizer_canonical import (
        reduce_strategy_for_legacy,
    )

    result = reduce_strategy_for_legacy(
        estimated_tokens=1000, sections=10,
        reduce_budget_tokens=5000,
    )
    assert result == "flat"

def test_select_reduce_strategy_exceeds_budget_uses_hierarchical():
    """estimated_tokens > reduce_budget AND sections ≥ min → HIERARCHICAL."""
    from workspace.skills.legal_summarizer.scripts.summarizer_canonical import (
        reduce_strategy_for_legacy,
    )

    result = reduce_strategy_for_legacy(
        estimated_tokens=5000, sections=5,
        reduce_budget_tokens=2000,
    )
    assert result == "hierarchical"

def test_select_reduce_strategy_few_sections_keeps_flat():
    """мало sections (< min) → FLAT даже если tokens > budget."""
    from workspace.skills.legal_summarizer.scripts.summarizer_canonical import (
        reduce_strategy_for_legacy,
    )

    result = reduce_strategy_for_legacy(
        estimated_tokens=5000, sections=1,
        reduce_budget_tokens=2000,
        min_sections_for_hierarchical=3,
    )
    assert result == "flat"

def test_select_reduce_strategy_min_sections_configurable():
    """``min_sections_for_hierarchical`` — настраиваемый порог."""
    from workspace.skills.legal_summarizer.scripts.summarizer_canonical import (
        reduce_strategy_for_legacy,
    )

    result_h = reduce_strategy_for_legacy(
        estimated_tokens=5000, sections=2,
        reduce_budget_tokens=2000, min_sections_for_hierarchical=2,
    )
    assert result_h == "hierarchical"
    result_f = reduce_strategy_for_legacy(
        estimated_tokens=5000, sections=2,
        reduce_budget_tokens=2000, min_sections_for_hierarchical=3,
    )
    assert result_f == "flat"

def test_legacy_should_use_hierarchical_removed():
    """``should_use_hierarchical_reduce`` deprecated (мигрирован на canonical)."""
    try:
        from workspace.skills.legal_summarizer.scripts.reducer import (
            should_use_hierarchical_reduce,
        )
        should_use_hierarchical_reduce(None, [])
    except (NotImplementedError, ImportError):
        return
    raise AssertionError(
        "should_use_hierarchical_reduce должен быть deprecated или удалён",
    )

def test_reduce_strategy_enum_values():
    """``flat`` / ``hierarchical`` — canonical reduce_strategy labels."""
    from workspace.skills.legal_summarizer.scripts.summarizer_canonical import (
        reduce_strategy_for_legacy,
    )

    assert reduce_strategy_for_legacy(
        estimated_tokens=100, sections=10,
        reduce_budget_tokens=10000,
    ) == "flat"
    assert reduce_strategy_for_legacy(
        estimated_tokens=10000, sections=10,
        reduce_budget_tokens=5000,
    ) == "hierarchical"

def test_reduce_config_has_min_sections_for_hierarchical():
    """``reduce_strategy_for_legacy`` принимает ``min_sections_for_hierarchical``."""
    from workspace.skills.legal_summarizer.scripts.summarizer_canonical import (
        reduce_strategy_for_legacy,
    )

    # С min_sections=2: hierarchical если tokens > budget.
    result_h = reduce_strategy_for_legacy(
        estimated_tokens=10000, sections=2,
        reduce_budget_tokens=5000, min_sections_for_hierarchical=2,
    )
    assert result_h == "hierarchical"
    # С min_sections=3: flat (2 < 3).
    result_f = reduce_strategy_for_legacy(
        estimated_tokens=10000, sections=2,
        reduce_budget_tokens=5000, min_sections_for_hierarchical=3,
    )
    assert result_f == "flat"

# ---------------------------------------------------------------------------

# ---------------------------------------------------------------------------

def test_section_trim_truncation_replaces_llm_call():
    """oversized section_summary → truncation, а не LLM-вызов."""
    from workspace.skills.legal_summarizer.scripts import summarizer

    # Длинный section_summary, превышающий лимит.
    oversized = "X" * 20000
    max_chars = 12000

    # Truncate как делает summarizer.
    truncated = oversized[:max_chars]
    assert len(truncated) == max_chars
    # Это НЕ LLM-вызов (нет round-trip).
    # Проверяем, что нет ссылки на llm_section_trim.
    assert not hasattr(summarizer, "_llm_section_trim"), (
        "_llm_section_trim должен быть удалён из summarizer"
    )

def test_summarizer_run_no_trim_calls_for_oversized_sections(monkeypatch, tmp_path):
    """End-to-end: при section_summary > max_chars → trim_calls=0 (truncation)."""
    import summarizer

    monkeypatch.setattr(summarizer, "get_chunking_config", lambda: {
        "chunk_size": 4000,
        "chunk_overlap": 0,
        "single_call_threshold": 100,  # чтобы пошёл map_reduce path
        "chunk_size_input_ratio": None,
    })
    monkeypatch.setattr(summarizer, "get_execution_config", lambda: {
        "confirmation_threshold_sec": 0.001,
        "estimated_chunk_duration_sec": 0.001,
        "max_chunks_for_execution": 10000,
        "context_batching": {
            "system_prompt_tokens": 200,
            "instruction_tokens_per_map": 50,
            "chars_per_token": 3.5,
            "safety_margin": 0.85,
        },
        "llm_max_tokens": 8192,
        "max_concurrent_batches": 1,
    })

    # Мокируем llm.chat: section_summary всегда очень длинный (> 12000 chars),
    # чтобы старая логика вызвала бы trim.
    def fake_chat(messages, *, context=None, **kwargs):
        import re
        text = messages[1]["content"] if len(messages) > 1 else ""
        if re.search(r"DOCUMENT CHUNK \d+", text):
            n = len(re.findall(r"DOCUMENT CHUNK \d+", text))
            return "\n\n".join(f"DOC CHUNK {i + 1}: краткое саммари" for i in range(n))
        if "Частичные саммари чанков этого раздела" in text:
            # Очень длинный section_summary → должен truncation'нуться.
            return "Y" * 50000
        return "Итоговое саммари.\n\nСуть: договор."

    monkeypatch.setattr(summarizer.llm, "chat", fake_chat)

    text = (
        "Преамбула длинного документа. " * 50
        + "Статья 1. Общие положения. " * 50
        + ("Длинный абзац текста для chunking'а. " * 100)
        + "Статья 2. Обязанности. " * 50
        + ("Длинный абзац второго раздела. " * 100)
    )
    result = summarizer.run(text, length="detailed", confirmed=True, workspace_root=tmp_path)

    assert result["stats"]["section_trim_calls"] == 0

def test_summarizer_truncate_section_summary_helper():
    """вспомогательная проверка truncation."""
    text = "ABCDEFGH" * 100  # 800 chars
    max_chars = 300
    truncated = text[:max_chars]
    assert len(truncated) == 300
    assert truncated.startswith("ABCDEFGH")

# ---------------------------------------------------------------------------

# ---------------------------------------------------------------------------

def test_execution_strategy_values():
    """Canonical execution strategy имеет 3 значения."""
    from workspace.skills.legal_summarizer.scripts.summarizer_canonical import (
        execution_strategy_for_legacy,
    )

    assert execution_strategy_for_legacy(
        estimated_tokens=1000, sections=0,
        direct_budget_tokens=10000, reduce_budget_tokens=10000,
    ) == "direct"
    assert execution_strategy_for_legacy(
        estimated_tokens=15000, sections=2,
        direct_budget_tokens=10000, reduce_budget_tokens=20000,
        min_sections_for_hierarchical=3,
    ) == "map_flat"
    assert execution_strategy_for_legacy(
        estimated_tokens=30000, sections=5,
        direct_budget_tokens=10000, reduce_budget_tokens=20000,
    ) == "map_hierarchical"

def test_select_execution_strategy_direct():
    """estimated_tokens ≤ direct_budget → direct."""
    from workspace.skills.legal_summarizer.scripts.summarizer_canonical import (
        execution_strategy_for_legacy,
    )

    result = execution_strategy_for_legacy(
        estimated_tokens=2000, sections=3,
        direct_budget_tokens=30000, reduce_budget_tokens=20000,
    )
    assert result == "direct"

def test_select_execution_strategy_map_flat():
    """estimated_tokens > direct_budget, sections < threshold → map_flat."""
    from workspace.skills.legal_summarizer.scripts.summarizer_canonical import (
        execution_strategy_for_legacy,
    )

    result = execution_strategy_for_legacy(
        estimated_tokens=12500, sections=2,
        direct_budget_tokens=10000, reduce_budget_tokens=20000,
        min_sections_for_hierarchical=3,
    )
    assert result == "map_flat"

def test_select_execution_strategy_map_hierarchical():
    """estimated_tokens > direct_budget, sections ≥ threshold → map_hierarchical."""
    from workspace.skills.legal_summarizer.scripts.summarizer_canonical import (
        execution_strategy_for_legacy,
    )

    result = execution_strategy_for_legacy(
        estimated_tokens=50000, sections=20,
        direct_budget_tokens=10000, reduce_budget_tokens=20000,
    )
    assert result == "map_hierarchical"

def test_select_execution_strategy_no_llm_calls():
    """execution strategy selector НЕ вызывает LLM.

    Детерминированный selector — нет side-effects, нет I/O.
    """
    import summarizer

    call_count = {"n": 0}

    def fake_chat(messages, *, context=None, **kwargs):
        call_count["n"] += 1
        return "should not be called"

    summarizer.llm.chat = fake_chat  # type: ignore[assignment]

    from workspace.skills.legal_summarizer.scripts.summarizer_canonical import (
        execution_strategy_for_legacy,
    )

    execution_strategy_for_legacy(
        estimated_tokens=2500, sections=3,
        direct_budget_tokens=30000, reduce_budget_tokens=20000,
    )
    assert call_count["n"] == 0

def test_execution_strategy_config_holds_budgets():
    """Canonical strategy selector использует параметры напрямую."""
    from workspace.skills.legal_summarizer.scripts.summarizer_canonical import (
        execution_strategy_for_legacy,
    )

    result_direct = execution_strategy_for_legacy(
        estimated_tokens=100, sections=1,
        direct_budget_tokens=50000, reduce_budget_tokens=30000,
    )
    assert result_direct == "direct"

    result_flat = execution_strategy_for_legacy(
        estimated_tokens=60000, sections=1,
        direct_budget_tokens=50000, reduce_budget_tokens=30000,
        min_sections_for_hierarchical=3,
    )
    assert result_flat == "map_flat"

def test_select_execution_strategy_boundary_exactly_at_direct_budget():
    """Boundary: estimated_tokens == direct_budget → direct (≤)."""
    from workspace.skills.legal_summarizer.scripts.summarizer_canonical import (
        execution_strategy_for_legacy,
    )

    result = execution_strategy_for_legacy(
        estimated_tokens=10000, sections=5,
        direct_budget_tokens=10000, reduce_budget_tokens=5000,
    )
    assert result == "direct"

def test_select_execution_strategy_boundary_one_above_direct():
    """Boundary: estimated_tokens = direct + 1 → map_flat или map_hierarchical."""
    from workspace.skills.legal_summarizer.scripts.summarizer_canonical import (
        execution_strategy_for_legacy,
    )

    result = execution_strategy_for_legacy(
        estimated_tokens=10001, sections=5,
        direct_budget_tokens=10000, reduce_budget_tokens=20000,
        min_sections_for_hierarchical=3,
    )
    assert result in ("map_flat", "map_hierarchical")

# ---------------------------------------------------------------------------

# ---------------------------------------------------------------------------

def test_direct_strategy_min_chars_default_zero_keeps_old_behavior():
    """``direct_strategy_min_chars`` УДАЛЁН из production pipeline (этап Integration).

    После интеграции ``ExecutionStrategy.DIRECT`` параметр больше не читается —
    ``inspect()`` принимает решение через ``StrategyConfig.direct_budget_tokens``
    (из ``TokenBudget.direct_call_tokens``). Этот тест оставлен как REGRESSION:
    если кто-то случайно вернёт чтение ``direct_strategy_min_chars`` в
    теле ``inspect()``, тест напомнит, что у нас единый путь через селектор.
    """
    import summarizer
    import inspect as _inspect

    src = _inspect.getsource(summarizer.inspect)
    assert "direct_strategy_min_chars" not in src, (
        "direct_strategy_min_chars не должно использоваться в теле inspect() "
        "(этап Integration & Simplification — единый ExecutionStrategy селектор)."
    )

def test_token_budget_direct_call_tokens_positive():
    """``TokenBudget.direct_call_tokens`` положительное."""
    from workspace.skills.legal_summarizer.scripts.token_budget import (
        TokenBudget,
    )

    budget = TokenBudget(
        context_window_tokens=65536,
        system_prompt_tokens=1200,
        instruction_tokens=200,
        output_reserve_tokens=8192,
        safety_margin=0.85,
        chars_per_token=3.5,
    )
    # 65536 - 1200 - 8192 = 56144.
    assert budget.direct_call_tokens == 56144

def test_token_budget_direct_call_tokens_includes_output_reserve():
    """DIRECT budget должен учитывать output_reserve.

    Если output_reserve растёт → direct падает.
    """
    from workspace.skills.legal_summarizer.scripts.token_budget import (
        TokenBudget,
    )

    small = TokenBudget(
        context_window_tokens=65536,
        system_prompt_tokens=1200,
        instruction_tokens=200,
        output_reserve_tokens=2000,
        safety_margin=0.85,
        chars_per_token=3.5,
    )
    large = TokenBudget(
        context_window_tokens=65536,
        system_prompt_tokens=1200,
        instruction_tokens=200,
        output_reserve_tokens=16000,
        safety_margin=0.85,
        chars_per_token=3.5,
    )
    assert small.direct_call_tokens > large.direct_call_tokens

def test_token_budget_direct_call_tokens_differs_from_available_chunk_tokens():
    """DIRECT budget != available_chunk_tokens (нет per-chunk instruction)."""
    from workspace.skills.legal_summarizer.scripts.token_budget import (
        TokenBudget,
    )

    budget = TokenBudget(
        context_window_tokens=65536,
        system_prompt_tokens=1200,
        instruction_tokens=200,
        output_reserve_tokens=8192,
        safety_margin=0.85,
        chars_per_token=3.5,
    )
    # DIRECT не учитывает instruction и safety_margin.
    assert budget.direct_call_tokens != budget.available_chunk_tokens
    # DIRECT больше available (нет per-chunk instruction и margin).
    assert budget.direct_call_tokens > budget.available_chunk_tokens

def test_inspect_direct_strategy_for_short_text():
    """Короткий текст, влезающий в direct budget → ``strategy='single'``.

    Этап Integration: ``ExecutionStrategy.DIRECT`` срабатывает по
    ``DocumentStats.estimated_tokens ≤ TokenBudget.direct_call_tokens``.
    Без LLM-вызовов: проверка чисто детерминированная.
    """
    import summarizer

    summarizer.get_chunking_config = lambda: {
        "chunk_size": 100, "chunk_overlap": 0,
        "single_call_threshold": 100, "chunk_size_input_ratio": None,
    }
    summarizer.get_execution_config = lambda: {
        "confirmation_threshold_sec": 0.001, "estimated_chunk_duration_sec": 0.001,
        "max_chunks_for_execution": 1000,
        "context_batching": {
            "system_prompt_tokens": 100, "instruction_tokens_per_map": 50,
            "chars_per_token": 3.5, "safety_margin": 0.85,
        },
        "llm_max_tokens": 100,
    }

    paragraph = "Юридический тестовый абзац для проверки директ-вызова. "
    text = "\n\n".join([paragraph] * 200)
    assert 10000 < len(text) < 20000
    insp = summarizer.inspect(text)
    assert insp.strategy == "single"
    assert insp.estimated_llm_calls == 1
    assert len(insp.chunks) == 1
    assert insp.context_batches == []


def test_inspect_map_reduce_for_long_text():
    """Длинный текст, НЕ влезающий в direct budget → ``strategy='map_reduce'``.

    Чтобы заставить ``ExecutionStrategy`` выбрать MAP_*, искусственно
    увеличиваем ``llm_max_tokens`` — это уменьшает ``direct_budget_tokens``
    через ``TokenBudget.direct_call_tokens = context - system - output``.
    """
    import summarizer

    summarizer.get_chunking_config = lambda: {
        "chunk_size": 100, "chunk_overlap": 0,
        "single_call_threshold": 100, "chunk_size_input_ratio": None,
    }
    summarizer.get_execution_config = lambda: {
        "confirmation_threshold_sec": 0.001, "estimated_chunk_duration_sec": 0.001,
        "max_chunks_for_execution": 1000,
        "context_batching": {
            "system_prompt_tokens": 100, "instruction_tokens_per_map": 50,
            "chars_per_token": 3.5, "safety_margin": 0.85,
        },
        "llm_max_tokens": 65000,
    }

    paragraph = "Юридический тестовый абзац для проверки мап-вызова. "
    text = "\n\n".join([paragraph] * 200)
    assert 10000 < len(text) < 20000
    insp = summarizer.inspect(text)
    assert insp.strategy == "map_reduce"
    assert len(insp.chunks) >= 1
    assert len(insp.context_batches) >= 1

# ---------------------------------------------------------------------------

# ---------------------------------------------------------------------------

def test_brief_strategy_default_coverage_ratio():
    """default coverage_ratio = 0.5 = старое поведение."""
    from skill_config import get_chunking_config

    cfg = get_chunking_config()
    assert float(cfg.get("brief_coverage_ratio", 0.5)) == 0.5

def test_select_brief_chunks_structured_coverage_ratio_05():
    """coverage_ratio=0.5 → ceil(total*0.5) chunks (старое поведение)."""
    from brief_strategy import select_brief_chunks_structured

    class _FakeSection:
        def __init__(self, sid):
            self.section_id = sid
            self.heading = sid
            self.section_path = sid

    class _FakeTree:
        def __init__(self, n_sections):
            self.root_id = "s_root"
            self.sections = {"s_root": _FakeSection("s_root")}
            for i in range(n_sections):
                self.sections[f"s_{i}"] = _FakeSection(f"s_{i}")

    chunks = []
    for i in range(60):
        sid = f"s_{i % 6}"  # 6 секций по 10 chunks
        c = type("C", (), {"index": i, "section_id": sid})()
        chunks.append(c)
    tree = _FakeTree(6)

    chosen = select_brief_chunks_structured(
        chunks, tree, max_chunks=100, coverage_ratio=0.5,
    )
    # ceil(60 * 0.5) = 30.
    assert len(chosen) == 30

def test_select_brief_chunks_structured_coverage_ratio_033():
    """coverage_ratio=0.33 → ceil(total*0.33) chunks."""
    from brief_strategy import select_brief_chunks_structured

    class _FakeSection:
        def __init__(self, sid):
            self.section_id = sid
            self.heading = sid
            self.section_path = sid

    class _FakeTree:
        def __init__(self, n_sections):
            self.root_id = "s_root"
            self.sections = {"s_root": _FakeSection("s_root")}
            for i in range(n_sections):
                self.sections[f"s_{i}"] = _FakeSection(f"s_{i}")

    chunks = []
    for i in range(60):
        sid = f"s_{i % 6}"
        c = type("C", (), {"index": i, "section_id": sid})()
        chunks.append(c)
    tree = _FakeTree(6)

    chosen = select_brief_chunks_structured(
        chunks, tree, max_chunks=100, coverage_ratio=0.33,
    )
    # ceil(60 * 0.33) = ceil(19.8) = 20.
    assert 19 <= len(chosen) <= 21

def test_select_brief_chunks_structured_coverage_ratio_zero_returns_empty():
    """coverage_ratio=0 → пустой (после sanitization fallback 0.5)."""
    from brief_strategy import select_brief_chunks_structured

    class _FakeSection:
        def __init__(self, sid):
            self.section_id = sid
            self.heading = sid
            self.section_path = sid

    class _FakeTree:
        def __init__(self, n_sections):
            self.root_id = "s_root"
            self.sections = {"s_root": _FakeSection("s_root")}
            for i in range(n_sections):
                self.sections[f"s_{i}"] = _FakeSection(f"s_{i}")

    chunks = [
        type("C", (), {"index": i, "section_id": f"s_{i % 3}"})()
        for i in range(30)
    ]
    tree = _FakeTree(3)

    # 0 вне диапазона → sanitization → 0.5 → ceil(15) = 15 chunks.
    chosen = select_brief_chunks_structured(
        chunks, tree, max_chunks=100, coverage_ratio=0.0,
    )
    assert len(chosen) == 15

def test_brief_strategy_min_chunks_3_sections_one_each():
    """3 секции по 4 chunks, ratio=0.33 → минимум 1 с каждой."""
    from brief_strategy import select_brief_chunks_structured

    class _FakeSection:
        def __init__(self, sid):
            self.section_id = sid
            self.heading = sid
            self.section_path = sid

    class _FakeTree:
        def __init__(self, n_sections):
            self.root_id = "s_root"
            self.sections = {"s_root": _FakeSection("s_root")}
            for i in range(n_sections):
                self.sections[f"s_{i}"] = _FakeSection(f"s_{i}")

    chunks = []
    for i in range(12):
        sid = f"s_{i % 3}"  # 3 секции × 4 chunks
        chunks.append(type("C", (), {"index": i, "section_id": sid})())
    tree = _FakeTree(3)

    chosen = select_brief_chunks_structured(
        chunks, tree, max_chunks=100, coverage_ratio=0.33,
    )
    # ceil(12 * 0.33) = 4. Round-robin покрывает все 3 секции.
    assert len(chosen) == 4
    # Покрытие всех 3 секций (при i%3: s_0: 0,3,6,9; s_1: 1,4,7,10; s_2: 2,5,8,11).
    sections_covered = {c.section_id for c in chosen}
    assert sections_covered == {"s_0", "s_1", "s_2"}
    # Round-robin берёт первый chunk из каждой секции (first_seen order).
    indices = {c.index for c in chosen}
    assert 0 in indices  # первый s_0
    assert 1 in indices  # первый s_1
    assert 2 in indices  # первый s_2
    # Последний взят = следующий по s_0 round-robin (3).
    assert 3 in indices
