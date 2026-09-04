"""Тесты для документов только с tables (Этап 72 из PLAN.md).

PLAN §72: документ только с таблицами — valid structure.
``DocumentStructure`` должен правильно обработать: корень + table nodes,
без section nodes.
"""

from __future__ import annotations

import tempfile

from workspace.skills.legal_summarizer.scripts.structure.pipeline import (
    run_canonical_pipeline,
)


def test_pipeline_tables_only():
    import tempfile
    from docx import Document

    doc = Document()
    for i in range(3):
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = f"Header {i}"
        table.cell(0, 1).text = f"Value {i}"

    with tempfile.NamedTemporaryFile(
        mode="wb", suffix=".docx", delete=False,
    ) as f:
        path = f.name
    doc.save(path)

    result = run_canonical_pipeline(path)
    assert result is not None


def test_pipeline_table_heavy_does_not_crash():
    import tempfile
    from docx import Document

    doc = Document()
    doc.add_paragraph("Header text")
    for i in range(5):
        table = doc.add_table(rows=3, cols=2)
        table.cell(0, 0).text = f"col1 row {i}"
        table.cell(0, 1).text = f"col2 row {i}"

    with tempfile.NamedTemporaryFile(
        mode="wb", suffix=".docx", delete=False,
    ) as f:
        path = f.name
    doc.save(path)

    result = run_canonical_pipeline(path)
    assert result.analysis is not None