"""PhysicalDocument: нормализованный список блоков документа с координатами.

Это **adapter** над ``workspace.utils.office_files``, не parser.

Что берём из office_files:
  * ``extract_structure(path)`` → ``title``, ``begin``, ``end``, ``text``,
    ``format``, ``size_bytes`` (full text).
  * ``extract_tables(path)`` → для DOCX/PDF: список таблиц.

Чего office_files не возвращает, и мы добавляем сами (точечные обёртки):
  * ``PdfReader.pages[i].extract_text()`` → отдельный ``DocumentBlock``
    с ``page_index=i+1``.
  * ``Document.paragraphs[i].text`` → отдельный ``DocumentBlock``
    с ``paragraph_index=i``.
  * ``Document.tables[i]`` → отдельный ``DocumentBlock``
    с ``table_index=i``.

Эти обёртки минимальны (буквально 1-2 строки на блок) и **не** пытаются
перепарсить документ — только сохраняют координаты, потерянные в
``extract_text``.
"""

from __future__ import annotations

import hashlib
import json
from dataclasses import asdict, dataclass, field
from pathlib import Path
from typing import Any

from workspace.utils.office_files import (
    detect_format,
    extract_tables,
)


SUPPORTED_FORMATS: frozenset[str] = frozenset({"pdf", "docx", "txt"})


@dataclass(frozen=True)
class DocumentBlock:
    """Один физический блок документа.

    Attributes:
        block_id: стабильный идентификатор вида ``"b_001"``.
        block_type: ``"page"`` (PDF) | ``"paragraph"`` (DOCX) |
            ``"table"`` | ``"text"`` (TXT) | ``"slide"`` (PPTX,
            зарезервировано на будущее).
        content: текст блока.
        char_count: ``len(content)``.
        page_index: 1-based номер страницы (PDF) или страница, на которой
            находится параграф/таблица (DOCX).
        page_start / page_end: для multi-page block (будущее). Сейчас
            всегда равны ``page_index``.
        paragraph_index: индекс параграфа в DOCX (None для не-DOCX).
        table_index: индекс таблицы в DOCX/PDF (None для не-таблицы).
        ordinal: 0..N-1, canonical document order (invariant #3).
        block_metadata: дополнительная мета (например, ``{"row_count": 5}``
            для таблицы).
    """

    block_id: str
    block_type: str
    content: str
    char_count: int
    page_index: int | None
    page_start: int | None
    page_end: int | None
    paragraph_index: int | None
    table_index: int | None
    ordinal: int
    block_metadata: dict[str, Any] = field(default_factory=dict)

    def to_dict(self) -> dict[str, Any]:
        return asdict(self)


@dataclass(frozen=True)
class PhysicalDocument:
    """Нормализованная физическая модель документа.

    Attributes:
        path: исходный путь.
        format: расширение файла (``pdf`` / ``docx`` / ``txt``).
        title: из ``extract_structure``.
        size_bytes: размер файла.
        blocks: плоский список ``DocumentBlock`` в canonical document order.
            ``blocks[i].ordinal == i`` (invariant #3).
        page_count: для PDF — число страниц; для DOCX — оценочное
            (max ``page_index`` среди blocks, не строгое).
    """

    path: str
    format: str
    title: str | None
    size_bytes: int
    blocks: tuple[DocumentBlock, ...]
    page_count: int

    def to_dict(self) -> dict[str, Any]:
        return {
            "path": self.path,
            "format": self.format,
            "title": self.title,
            "size_bytes": self.size_bytes,
            "page_count": self.page_count,
            "blocks": [b.to_dict() for b in self.blocks],
        }

    @classmethod
    def from_dict(cls, data: dict[str, Any]) -> "PhysicalDocument":
        blocks = tuple(DocumentBlock(**b) for b in data["blocks"])
        return cls(
            path=data["path"],
            format=data["format"],
            title=data.get("title"),
            size_bytes=data["size_bytes"],
            blocks=blocks,
            page_count=data["page_count"],
        )


def _physical_cache_root(workspace_root: Path | str | None) -> Path:
    if workspace_root is None:
        return Path("workspace/data_store/cache/skills/legal_summarizer/physical")
    return Path(workspace_root) / "workspace" / "data_store" / "cache" / "skills" / "legal_summarizer" / "physical"


def _physical_cache_key(path: Path) -> str:
    st = path.stat()
    raw = f"{path.resolve()}|{st.st_size}|{st.st_mtime}"
    return hashlib.sha256(raw.encode("utf-8")).hexdigest()[:32]


def _pick_title_from_text(fmt: str, text: str, path: Path) -> str | None:
    """Получить title из метаданных файла или первой содержательной строки."""
    if fmt == "docx":
        try:
            from docx import Document

            doc = Document(str(path))
            if doc.core_properties.title:
                return doc.core_properties.title.strip()
        except Exception:
            pass
    if fmt == "pdf":
        try:
            from pypdf import PdfReader

            reader = PdfReader(str(path))
            title = (reader.metadata or {}).get("/Title")
            if title:
                return str(title).strip()
        except Exception:
            pass
    if fmt == "pptx":
        try:
            from pptx import Presentation

            pres = Presentation(str(path))
            if pres.core_properties.title:
                return pres.core_properties.title.strip()
        except Exception:
            pass
    for line in (text or "").splitlines():
        line = line.strip()
        if len(line) >= 3:
            return line[:200]
    return None


def _build_structure_dict(path: Path, fmt: str) -> dict[str, Any]:
    """Построить минимальный dict с полями, нужными PhysicalDocument.

    Поля: ``title``, ``text``, ``begin``, ``end``, ``format``, ``size_bytes``.
    Аналог ранее существовавшего ``office_files.extract_structure``.
    """
    text = ""
    text_error: str | None = None
    try:
        from workspace.utils.office_files import extract_text
        text = extract_text(path)
    except Exception as e:
        text_error = str(e)

    title = _pick_title_from_text(fmt, text, path)
    size_bytes = path.stat().st_size
    info: dict[str, Any] = {
        "title": title,
        "format": fmt,
        "size_bytes": size_bytes,
    }
    if text_error:
        info["text_error"] = text_error
    return info


def _table_to_text(table: list[list[str]]) -> str:
    """Склеить таблицу в текст с ``|``-разделителем ячеек (как office_files)."""
    lines: list[str] = []
    for row in table:
        cells = [c.strip() if c else "" for c in row]
        if any(cells):
            lines.append(" | ".join(cells))
    return "\n".join(lines)


def _iter_pdf_blocks(path: Path) -> tuple[list[DocumentBlock], int]:
    """PDF → blocks (по страницам). Таблицы встроены между страницами."""
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    blocks: list[DocumentBlock] = []
    ordinal = 0
    page_count = len(reader.pages)

    page_to_tables: dict[int, list[list[list[str]]]] = {}
    try:
        import pdfplumber

        with pdfplumber.open(str(path)) as pdf:
            for p_idx, page in enumerate(pdf.pages, start=1):
                tables = page.extract_tables() or []
                if tables:
                    page_to_tables[p_idx] = [
                        [[cell or "" for cell in row] for row in t] for t in tables
                    ]
    except Exception:
        page_to_tables = {}

    for page_idx, page in enumerate(reader.pages, start=1):
        try:
            text = page.extract_text() or ""
        except Exception:
            text = ""
        if text.strip():
            blocks.append(
                DocumentBlock(
                    block_id=f"b_{ordinal:04d}",
                    block_type="page",
                    content=text,
                    char_count=len(text),
                    page_index=page_idx,
                    page_start=page_idx,
                    page_end=page_idx,
                    paragraph_index=None,
                    table_index=None,
                    ordinal=ordinal,
                    block_metadata={},
                )
            )
            ordinal += 1

        for t_idx, table in enumerate(page_to_tables.get(page_idx, [])):
            table_text = _table_to_text(table)
            if not table_text.strip():
                continue
            blocks.append(
                DocumentBlock(
                    block_id=f"b_{ordinal:04d}",
                    block_type="table",
                    content=table_text,
                    char_count=len(table_text),
                    page_index=page_idx,
                    page_start=page_idx,
                    page_end=page_idx,
                    paragraph_index=None,
                    table_index=t_idx,
                    ordinal=ordinal,
                    block_metadata={"row_count": len(table)},
                )
            )
            ordinal += 1

    return blocks, page_count


def _iter_docx_blocks(path: Path) -> tuple[list[DocumentBlock], int]:
    """DOCX → blocks (paragraphs + tables) в document order."""
    from docx import Document

    doc = Document(str(path))
    blocks: list[DocumentBlock] = []
    ordinal = 0

    para_to_page: dict[int, int] = {}
    page_no = 1
    for p_idx in range(len(doc.paragraphs)):
        if p_idx > 0 and p_idx % 25 == 0:
            page_no += 1
        para_to_page[p_idx] = page_no

    for p_idx, para in enumerate(doc.paragraphs):
        text = (para.text or "").strip()
        if not text:
            continue
        page_idx = para_to_page.get(p_idx)
        blocks.append(
            DocumentBlock(
                block_id=f"b_{ordinal:04d}",
                block_type="paragraph",
                content=text,
                char_count=len(text),
                page_index=page_idx,
                page_start=page_idx,
                page_end=page_idx,
                paragraph_index=p_idx,
                table_index=None,
                ordinal=ordinal,
                block_metadata={"style": para.style.name if para.style else ""},
            )
        )
        ordinal += 1

    for t_idx, table in enumerate(doc.tables):
        rows: list[list[str]] = []
        for row in table.rows:
            rows.append([cell.text.strip() for cell in row.cells])
        table_text = _table_to_text(rows)
        if not table_text.strip():
            continue
        blocks.append(
            DocumentBlock(
                block_id=f"b_{ordinal:04d}",
                block_type="table",
                content=table_text,
                char_count=len(table_text),
                page_index=None,
                page_start=None,
                page_end=None,
                paragraph_index=None,
                table_index=t_idx,
                ordinal=ordinal,
                block_metadata={"row_count": len(rows)},
            )
        )
        ordinal += 1

    page_count = max((b.page_index or 0) for b in blocks) if blocks else 1
    return blocks, page_count


def _iter_txt_blocks(path: Path) -> tuple[list[DocumentBlock], int]:
    """TXT → один block."""
    raw = path.read_bytes()
    try:
        text = raw.decode("utf-8")
    except UnicodeDecodeError:
        text = raw.decode("utf-8", errors="replace")
    blocks = [
        DocumentBlock(
            block_id="b_0000",
            block_type="text",
            content=text,
            char_count=len(text),
            page_index=None,
            page_start=None,
            page_end=None,
            paragraph_index=None,
            table_index=None,
            ordinal=0,
            block_metadata={},
        )
    ]
    return blocks, 1


def load_physical_document(
    path: str | Path,
    *,
    workspace_root: Path | str | None = None,
) -> PhysicalDocument:
    """Загрузить PhysicalDocument для файла с кэшированием на диске.

    Args:
        path: путь к PDF/DOCX/TXT.
        workspace_root: корень workspace (для тестов).

    Returns:
        PhysicalDocument с нормализованным списком DocumentBlock.

    Raises:
        FileNotFoundError: файл не существует.
        ValueError: формат не поддерживается.
    """
    p = Path(path)
    if not p.exists():
        raise FileNotFoundError(f"Файл не найден: {p}")

    fmt = detect_format(p)
    if fmt not in SUPPORTED_FORMATS:
        raise ValueError(
            f"Неподдерживаемый формат для PhysicalDocument: '{fmt}'. "
            f"Поддерживаются: {sorted(SUPPORTED_FORMATS)}"
        )

    cache_dir = _physical_cache_root(workspace_root)
    cache_file = cache_dir / f"{_physical_cache_key(p)}.json"

    if cache_file.is_file():
        try:
            cached = json.loads(cache_file.read_text(encoding="utf-8"))
            if cached.get("path") == str(p.resolve()) and cached.get("format") == fmt:
                return PhysicalDocument.from_dict(cached)
        except (OSError, json.JSONDecodeError, KeyError, TypeError):
            pass

    struct = _build_structure_dict(p, fmt)
    title = struct.get("title")
    size_bytes = struct.get("size_bytes", p.stat().st_size)

    if fmt == "pdf":
        blocks, page_count = _iter_pdf_blocks(p)
    elif fmt == "docx":
        blocks, page_count = _iter_docx_blocks(p)
    elif fmt == "txt":
        blocks, page_count = _iter_txt_blocks(p)
    else:
        raise ValueError(f"Unsupported format: {fmt}")

    doc = PhysicalDocument(
        path=str(p.resolve()),
        format=fmt,
        title=title,
        size_bytes=size_bytes,
        blocks=tuple(blocks),
        page_count=page_count,
    )

    try:
        cache_dir.mkdir(parents=True, exist_ok=True)
        cache_file.write_text(
            json.dumps(doc.to_dict(), ensure_ascii=False),
            encoding="utf-8",
        )
    except OSError:
        pass

    return doc


__all__ = [
    "DocumentBlock",
    "PhysicalDocument",
    "load_physical_document",
    "SUPPORTED_FORMATS",
]